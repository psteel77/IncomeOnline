"""
Build PDF versions of the MoneyRules guides from their .docx sources.

The .docx files in /static are the editable masters (the user maintains them in
Montserrat). This tool produces print-ready PDFs that render identically on
every device, with the Montserrat font fully embedded.

It is a BUILD-TIME tool: it needs LibreOffice (`soffice`) and the Montserrat
font installed. The generated PDFs are committed to /static and served directly
in production (Railway needs no LibreOffice at runtime).

Run:  python build_guide_pdfs.py
"""
import os
import subprocess
import tempfile
import shutil

from docx import Document
from docx.oxml.ns import qn

STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")

# The 10 free guides (+ blank template) — docx master -> pdf output (same base name).
GUIDE_DOCX = [
    "The_Rule_of_72_Guide.docx",
    "The_50_30_20_Budget_Rule.docx",
    "Passive_Income_Beginners_Guide.docx",
    "The_Debt_Snowball_Method.docx",
    "The_Emergency_Fund_Guide.docx",
    "Compound_Interest_Handbook.docx",
    "UK_Tax_Basics_Freelancers.docx",
    "UK_Credit_Score_Masterclass.docx",
    "ISA_vs_SIPP_Complete_Guide.docx",
    "Side_Hustle_Quick_Start_Guide.docx",
]

FONT = "Montserrat"


def _set_run_font(run):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def _set_style_font(style):
    try:
        style.font.name = FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), FONT)
    except Exception:
        pass


def _normalize_fonts(doc):
    """Force every style and run to Montserrat (font family only — no layout change)."""
    for style in doc.styles:
        _set_style_font(style)

    def walk_paragraphs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                _set_run_font(r)

    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    walk_paragraphs(cell.paragraphs)
                    walk_tables(cell.tables)

    walk_paragraphs(doc.paragraphs)
    walk_tables(doc.tables)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            try:
                walk_paragraphs(hf.paragraphs)
                walk_tables(hf.tables)
            except Exception:
                pass


def docx_to_pdf(docx_path, out_dir):
    """Normalize fonts -> convert to PDF via LibreOffice. Returns the PDF path."""
    base = os.path.splitext(os.path.basename(docx_path))[0]
    doc = Document(docx_path)
    _normalize_fonts(doc)

    with tempfile.TemporaryDirectory() as tmp:
        norm_path = os.path.join(tmp, f"{base}.docx")
        doc.save(norm_path)
        env = dict(os.environ, HOME=tmp)
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, norm_path],
            check=True, capture_output=True, timeout=120, env=env,
        )
        produced = os.path.join(tmp, f"{base}.pdf")
        final = os.path.join(out_dir, f"{base}.pdf")
        shutil.move(produced, final)
        return final


def buffer_to_pdf(buffer, out_path):
    """Normalize fonts in an in-memory .docx buffer -> convert to PDF at out_path."""
    base = os.path.splitext(os.path.basename(out_path))[0]
    doc = Document(buffer)
    _normalize_fonts(doc)
    with tempfile.TemporaryDirectory() as tmp:
        norm_path = os.path.join(tmp, f"{base}.docx")
        doc.save(norm_path)
        env = dict(os.environ, HOME=tmp)
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, norm_path],
            check=True, capture_output=True, timeout=120, env=env,
        )
        produced = os.path.join(tmp, f"{base}.pdf")
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        shutil.move(produced, out_path)
        return out_path


def build_premium_pdfs():
    """Build the Premium Pack's welcome + 2 exclusive guide PDFs into static/premium/."""
    from generate_premium_pack import (
        _readme_docx, _premium_guide_wealth_roadmap, _premium_guide_fire_playbook,
    )
    premium_dir = os.path.join(STATIC_DIR, "premium")
    targets = [
        (_readme_docx, "00_WELCOME_START_HERE.pdf"),
        (_premium_guide_wealth_roadmap, "11_Wealth_Building_Roadmap_PREMIUM.pdf"),
        (_premium_guide_fire_playbook, "12_The_FIRE_Playbook_PREMIUM.pdf"),
    ]
    built = []
    for gen_fn, out_name in targets:
        out_path = buffer_to_pdf(gen_fn(), os.path.join(premium_dir, out_name))
        print(f"  built premium/{out_name} ({os.path.getsize(out_path)} bytes)")
        built.append(out_path)
    return built


def build_all():
    built = []
    for name in GUIDE_DOCX:
        src = os.path.join(STATIC_DIR, name)
        if not os.path.exists(src):
            print(f"  SKIP (missing): {name}")
            continue
        pdf = docx_to_pdf(src, STATIC_DIR)
        print(f"  built {os.path.basename(pdf)} ({os.path.getsize(pdf)} bytes)")
        built.append(pdf)
    return built


if __name__ == "__main__":
    print(f"Building guide PDFs into {STATIC_DIR} ...")
    build_all()
    print("Building premium pack PDFs ...")
    build_premium_pdfs()
    print("Building premium pack ZIP ...")
    from generate_premium_pack import build_premium_pack
    build_premium_pack()
    print("Done.")
