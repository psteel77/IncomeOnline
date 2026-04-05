"""
Generate The Rule of 72 guide using the MoneyRules branded template.
"""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, DEEP_PURPLE, PURPLE, PINK, ACCENT_GOLD, DARK_TEXT, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


def generate_rule72_document():
    """Generate the complete Rule of 72 Word document using MoneyRules template."""
    doc = create_moneyrules_document(
        title='The Rule of 72',
        subtitle='A Complete Guide to Estimating Investment Growth'
    )

    # ===== TITLE PAGE =====
    add_title_page(
        doc,
        title='The Rule of 72',
        subtitle='A Complete Guide to Estimating Investment Growth',
        tagline='How to quickly calculate when your money will double\n— and why every investor should know this formula'
    )

    # ===== TABLE OF CONTENTS =====
    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()

    toc_items = [
        ('1.', 'Introduction — What Is the Rule of 72?'),
        ('2.', 'The Formula Explained'),
        ('3.', 'Step-by-Step: How to Use the Rule of 72'),
        ('4.', 'Quick Reference Table — Doubling Times'),
        ('5.', 'Real-World Investment Examples'),
        ('6.', 'The Power of Compound Interest'),
        ('7.', 'Accuracy Check: Rule of 72 vs. Exact Calculations'),
        ('8.', 'Variations: The Rule of 69 and Rule of 70'),
        ('9.', 'Limitations and When Not to Use It'),
        ('10.', 'Key Takeaways and Conclusion'),
    ]

    for num, title_text in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(10)
        run_num = para.add_run(f'{num}  ')
        run_num.font.size = Pt(13)
        run_num.font.color.rgb = PURPLE
        run_num.bold = True
        run_title = para.add_run(title_text)
        run_title.font.size = Pt(13)
        run_title.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run(
        'DISCLAIMER: This document is for educational and informational purposes only. '
        'It does not constitute financial advice. Always consult a qualified financial '
        'adviser before making investment decisions.'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    run.italic = True
    doc.add_page_break()

    # ===== CHAPTER 1: INTRODUCTION =====
    add_styled_heading(doc, '1. Introduction — What Is the Rule of 72?', level=1)

    add_body_text(doc,
        'The Rule of 72 is one of the most useful shortcuts in personal finance and investing. '
        'It allows you to quickly estimate how long it will take for an investment to double in value, '
        'given a fixed annual rate of return — all without needing a calculator, spreadsheet, or '
        'complex mathematical formula.')

    add_body_text(doc,
        'The concept is beautifully simple: divide the number 72 by your expected annual interest rate '
        '(expressed as a whole number), and the result gives you the approximate number of years '
        'it will take for your money to double.')

    add_highlight_box(doc, 'Years to Double  =  72  ÷  Annual Interest Rate (%)')

    add_body_text(doc,
        'For example, if you invest money at an annual return of 8%, it will take approximately '
        '72 ÷ 8 = 9 years for your investment to double. If you earn 6% per year, your money '
        'doubles in about 72 ÷ 6 = 12 years.')

    add_styled_heading(doc, 'A Brief History', level=2)
    add_body_text(doc,
        'The Rule of 72 dates back to at least 1494, when the Italian mathematician Luca Pacioli '
        'referenced it in his work "Summa de Arithmetica." It has since become a staple of financial '
        'education worldwide, taught in universities, used by financial advisers, and relied upon by '
        'everyday investors who want a quick mental shortcut for understanding growth.')

    add_styled_heading(doc, 'Why Is It Called the "Rule of 72"?', level=2)
    add_body_text(doc,
        'The number 72 is used because it is a convenient approximation that works well across a wide '
        'range of realistic interest rates (roughly 2% to 20%). It also has many divisors (1, 2, 3, 4, '
        '6, 8, 9, 12, 18, 24, 36, 72), making mental arithmetic easy. The true mathematical constant '
        'would be closer to 69.3 (the natural logarithm of 2 times 100), but 72 provides a better '
        'practical approximation for typical investment rates and is far easier to divide in your head.')

    add_body_text(doc,
        'Whether you are a seasoned investor managing a portfolio or a beginner just starting to save, '
        'the Rule of 72 is an invaluable tool that you will use again and again throughout your '
        'financial journey.')
    doc.add_page_break()

    # ===== CHAPTER 2: FORMULA =====
    add_styled_heading(doc, '2. The Formula Explained', level=1)

    add_body_text(doc,
        'At its core, the Rule of 72 is a simplified version of the compound interest formula. '
        'Let us break down both so you can see where this powerful shortcut comes from.')

    add_styled_heading(doc, 'The Compound Interest Formula', level=2)
    add_body_text(doc, 'The standard formula for compound interest is:')

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(12)
    formula.paragraph_format.space_after = Pt(12)
    run = formula.add_run('A = P × (1 + r)ⁿ')
    run.font.size = Pt(16)
    run.font.color.rgb = DEEP_PURPLE
    run.bold = True

    add_body_text(doc, 'Where:')
    for item in [
        'A = the future value of the investment',
        'P = the principal (initial investment)',
        'r = the annual interest rate (as a decimal, e.g. 0.08 for 8%)',
        'n = the number of years'
    ]:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    add_body_text(doc, 'To find when the investment doubles, we set A = 2P:')

    formula2 = doc.add_paragraph()
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula2.add_run('2P = P × (1 + r)ⁿ  →  n = ln(2) / ln(1 + r)')
    run.font.size = Pt(14)
    run.font.color.rgb = PINK
    run.bold = True

    add_body_text(doc,
        'For small values of r, ln(1 + r) is approximately equal to r. Since ln(2) = 0.693, '
        'we get n ≈ 0.693 / r, or equivalently, n ≈ 69.3 / R (where R is the rate as a percentage). '
        'Rounding 69.3 up to 72 provides a better approximation across a broader range of rates and '
        'is much easier to compute mentally. This is where "72" comes from.')

    add_highlight_box(doc, 'The Rule of 72 is a mental-math-friendly approximation\nof the exact compound interest doubling formula.')

    add_styled_heading(doc, 'Key Assumptions', level=2)
    for a in [
        'The interest rate is fixed (does not change year to year).',
        'Returns are compounded annually (not monthly, daily, or continuously).',
        'No additional deposits or withdrawals are made.',
        'Taxes and fees are not taken into account.'
    ]:
        para = doc.add_paragraph(a, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    add_body_text(doc,
        'Despite these simplifications, the Rule of 72 remains remarkably accurate for interest rates '
        'between 2% and 20%, which covers the vast majority of real-world investment scenarios.')
    doc.add_page_break()

    # ===== CHAPTER 3: HOW TO USE =====
    add_styled_heading(doc, '3. Step-by-Step: How to Use the Rule of 72', level=1)

    add_body_text(doc, 'Using the Rule of 72 is straightforward. Here is a simple three-step process:')

    add_styled_heading(doc, 'Step 1: Identify Your Annual Rate of Return', level=2)
    add_body_text(doc,
        'Determine the annual interest rate or expected rate of return on your investment. '
        'This could be the interest rate on a savings account, the average annual return of '
        'a stock index fund, or the yield on a bond. Express it as a whole number (e.g., 6 for 6%).')

    add_styled_heading(doc, 'Step 2: Divide 72 by That Rate', level=2)
    add_body_text(doc,
        'Simply divide 72 by the interest rate. The result is the approximate number of years '
        'it takes for your investment to double.')

    add_styled_heading(doc, 'Step 3: Interpret Your Result', level=2)
    add_body_text(doc,
        'The answer tells you how many years until your money doubles at that rate of return, '
        'assuming the rate stays constant and returns are reinvested.')

    doc.add_paragraph()
    add_styled_heading(doc, 'Worked Examples', level=2)

    add_body_text(doc, 'Example 1: Savings Account at 3% Interest', bold=True)
    add_body_text(doc, 'You deposit £10,000 into a savings account paying 3% per year.')
    add_body_text(doc, 'Calculation: 72 ÷ 3 = 24 years')
    add_body_text(doc, 'Result: Your £10,000 will grow to approximately £20,000 in 24 years.')
    doc.add_paragraph()

    add_body_text(doc, 'Example 2: Stock Market Index Fund at 8% Return', bold=True)
    add_body_text(doc, 'You invest £5,000 in an index fund that averages 8% annual return.')
    add_body_text(doc, 'Calculation: 72 ÷ 8 = 9 years')
    add_body_text(doc, 'Result: Your £5,000 will grow to approximately £10,000 in 9 years.')
    doc.add_paragraph()

    add_body_text(doc, 'Example 3: High-Growth Investment at 12% Return', bold=True)
    add_body_text(doc, 'You invest £2,000 in a growth fund averaging 12% per year.')
    add_body_text(doc, 'Calculation: 72 ÷ 12 = 6 years')
    add_body_text(doc, 'Result: Your £2,000 will grow to approximately £4,000 in just 6 years.')

    add_highlight_box(doc, 'The higher the rate of return, the faster your money doubles.\nThis is the power of compound interest at work.')
    doc.add_page_break()

    # ===== CHAPTER 4: REFERENCE TABLE =====
    add_styled_heading(doc, '4. Quick Reference Table — Doubling Times', level=1)

    add_body_text(doc,
        'The table below shows how long it takes for an investment to double at various annual rates '
        'of return, using the Rule of 72. Keep this table handy as a quick reference whenever you '
        'are evaluating investment opportunities.')
    doc.add_paragraph()

    add_branded_table(doc,
        headers=['Annual Rate\nof Return', 'Years to\nDouble', 'Example: £1,000\nBecomes...', 'Typical Investment\nType'],
        data=[
            ('2%',  '36 years',   '£2,000 in 36 yrs',  'Government Bonds'),
            ('3%',  '24 years',   '£2,000 in 24 yrs',  'High-Interest Savings'),
            ('4%',  '18 years',   '£2,000 in 18 yrs',  'Corporate Bonds'),
            ('5%',  '14.4 years', '£2,000 in 14.4 yrs', 'Balanced Fund'),
            ('6%',  '12 years',   '£2,000 in 12 yrs',  'Dividend Stocks'),
            ('7%',  '10.3 years', '£2,000 in 10.3 yrs', 'Global Index Fund'),
            ('8%',  '9 years',    '£2,000 in 9 yrs',   'S&P 500 Average'),
            ('9%',  '8 years',    '£2,000 in 8 yrs',   'Growth Stocks'),
            ('10%', '7.2 years',  '£2,000 in 7.2 yrs', 'Aggressive Growth Fund'),
            ('12%', '6 years',    '£2,000 in 6 yrs',   'Emerging Markets'),
            ('15%', '4.8 years',  '£2,000 in 4.8 yrs', 'Venture / High Risk'),
            ('18%', '4 years',    '£2,000 in 4 yrs',   'Exceptional Returns'),
        ]
    )

    add_body_text(doc,
        'As you can see, even small differences in return rates have a dramatic impact on how '
        'quickly wealth accumulates. An investment returning 8% doubles in 9 years, while one '
        'returning 4% takes 18 years — twice as long.')

    add_highlight_box(doc, 'A 1% increase in annual return can shave years\noff the time it takes to double your money.')
    doc.add_page_break()

    # ===== CHAPTER 5: REAL-WORLD EXAMPLES =====
    add_styled_heading(doc, '5. Real-World Investment Examples', level=1)

    add_body_text(doc, 'Let us apply the Rule of 72 to some real-world scenarios.')

    add_styled_heading(doc, 'Scenario 1: Retirement Planning', level=2)
    add_body_text(doc,
        'Sarah is 25 years old and invests £10,000 in a diversified stock index fund that '
        'historically returns about 7% per year. She wants to know how her investment will grow '
        'by the time she retires at 65 — that is 40 years away.')
    add_body_text(doc, 'Using the Rule of 72: 72 ÷ 7 = approximately 10.3 years to double.')
    add_body_text(doc, 'In 40 years, her money will double roughly 4 times (40 ÷ 10.3 ≈ 3.9):')

    add_branded_table(doc,
        headers=['Stage', 'Value', 'Approximate Age'],
        data=[
            ('Start',       '£10,000',  'Age 25'),
            ('1st Doubling', '£20,000', 'Age 35'),
            ('2nd Doubling', '£40,000', 'Age 45'),
            ('3rd Doubling', '£80,000', 'Age 55'),
            ('4th Doubling', '£160,000', 'Age 65'),
        ],
        header_color='DB2777'
    )

    add_body_text(doc,
        'Sarah\'s initial £10,000 could grow to approximately £160,000 by retirement — a 16x '
        'increase — simply by investing early and allowing compound interest to work over time.',
        italic=True)

    add_styled_heading(doc, 'Scenario 2: Comparing Two Investments', level=2)
    add_body_text(doc,
        'James is considering two investment options:\n'
        '   Option A: A bond fund returning 4% per year\n'
        '   Option B: A stock index fund returning 8% per year')
    add_body_text(doc,
        'Using the Rule of 72:\n'
        '   Option A: 72 ÷ 4 = 18 years to double\n'
        '   Option B: 72 ÷ 8 = 9 years to double')
    add_body_text(doc,
        'If James invests £20,000 for 36 years:\n'
        '   Option A doubles twice (36 ÷ 18 = 2):  £20,000 → £40,000 → £80,000\n'
        '   Option B doubles four times (36 ÷ 9 = 4):  £20,000 → £40,000 → £80,000 → £160,000 → £320,000')

    add_highlight_box(doc, 'The difference between 4% and 8% over 36 years:\n£80,000 vs £320,000 — a fourfold difference!')

    add_styled_heading(doc, 'Scenario 3: The Cost of Inflation', level=2)
    add_body_text(doc,
        'The Rule of 72 also works in reverse — you can use it to understand how quickly inflation '
        'erodes your purchasing power. If inflation averages 3% per year:')
    add_body_text(doc, 'Calculation: 72 ÷ 3 = 24 years')
    add_body_text(doc,
        'This means the purchasing power of your money is cut in half every 24 years. £100 today '
        'will only buy £50 worth of goods in 24 years. This is a powerful reminder of why keeping '
        'your money in a zero-interest account is actually losing you money over time.')
    doc.add_page_break()

    # ===== CHAPTER 6: COMPOUND INTEREST =====
    add_styled_heading(doc, '6. The Power of Compound Interest', level=1)

    add_body_text(doc,
        'Albert Einstein is often quoted as saying, "Compound interest is the eighth wonder of '
        'the world. He who understands it, earns it; he who doesn\'t, pays it." While the '
        'attribution is debated, the truth of the statement is not.')

    add_body_text(doc,
        'The Rule of 72 gives us a window into the exponential nature of compound interest. '
        'Unlike simple interest (where you only earn returns on your original investment), '
        'compound interest means you earn returns on your returns.')

    add_styled_heading(doc, 'Simple Interest vs Compound Interest', level=2)
    add_body_text(doc, 'Consider £10,000 invested at 8% for 30 years:')

    add_branded_table(doc,
        headers=['Type', 'Calculation', 'Final Value'],
        data=[
            ('Simple Interest', '£10,000 + (£800 × 30)', '£34,000'),
            ('Compound Interest', '£10,000 × (1.08)³⁰', '£100,627'),
        ],
        header_color='EA580C'
    )

    add_body_text(doc,
        'The difference is staggering: £34,000 with simple interest versus over £100,000 with '
        'compound interest. That is nearly three times as much money, all because the returns '
        'themselves earn returns year after year.')

    add_styled_heading(doc, 'Multiple Doublings Over Time', level=2)
    add_body_text(doc,
        'The Rule of 72 helps us visualise this compounding effect through successive doublings:')

    for p in [
        '1st doubling: £1,000 → £2,000  (gain of £1,000)',
        '2nd doubling: £2,000 → £4,000  (gain of £2,000)',
        '3rd doubling: £4,000 → £8,000  (gain of £4,000)',
        '4th doubling: £8,000 → £16,000  (gain of £8,000)',
        '5th doubling: £16,000 → £32,000  (gain of £16,000)',
        '6th doubling: £32,000 → £64,000  (gain of £32,000)',
    ]:
        para = doc.add_paragraph(p, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    add_body_text(doc,
        'Notice how the gains in each doubling period are equal to the total of ALL previous gains '
        'combined. This is the exponential magic of compounding.')

    add_highlight_box(doc, 'Time is the most powerful factor in wealth building.\nThe earlier you start investing, the more doublings you achieve.')
    doc.add_page_break()

    # ===== CHAPTER 7: ACCURACY CHECK =====
    add_styled_heading(doc, '7. Accuracy Check: Rule of 72 vs. Exact Calculations', level=1)

    add_body_text(doc,
        'How accurate is the Rule of 72? Let us compare its estimates with the exact mathematical '
        'answer (calculated using n = ln(2) / ln(1 + r)).')
    doc.add_paragraph()

    add_branded_table(doc,
        headers=['Interest Rate', 'Rule of 72\nEstimate', 'Exact\nAnswer', 'Difference'],
        data=[
            ('2%',  '36.00', '35.00', '+1.00 years'),
            ('3%',  '24.00', '23.45', '+0.55 years'),
            ('4%',  '18.00', '17.67', '+0.33 years'),
            ('5%',  '14.40', '14.21', '+0.19 years'),
            ('6%',  '12.00', '11.90', '+0.10 years'),
            ('7%',  '10.29', '10.24', '+0.05 years'),
            ('8%',  '9.00',  '9.01',  '-0.01 years'),
            ('9%',  '8.00',  '8.04',  '-0.04 years'),
            ('10%', '7.20',  '7.27',  '-0.07 years'),
            ('12%', '6.00',  '6.12',  '-0.12 years'),
        ]
    )

    add_body_text(doc,
        'The Rule of 72 is remarkably accurate for interest rates between 6% and 10%, where the '
        'error is less than a tenth of a year. Even at extreme rates, the maximum error is only '
        'about 1 year — perfectly acceptable for a quick mental estimate.')

    add_body_text(doc,
        'The Rule of 72 is most precise around 8%, where it is almost exactly correct. '
        'At lower rates (2-4%), it slightly overestimates the doubling time. At higher rates '
        '(above 10%), it slightly underestimates.')

    add_highlight_box(doc, 'The Rule of 72 is accurate to within 1 year\nfor interest rates between 2% and 18%.')
    doc.add_page_break()

    # ===== CHAPTER 8: VARIATIONS =====
    add_styled_heading(doc, '8. Variations: The Rule of 69 and Rule of 70', level=1)

    add_body_text(doc, 'While the Rule of 72 is the most popular, there are two common variations:')

    add_styled_heading(doc, 'The Rule of 69.3 (or Rule of 69)', level=2)
    add_body_text(doc,
        'Since ln(2) = 0.6931, dividing 69.3 by the interest rate gives the mathematically exact '
        'answer for continuous compounding. This rule is preferred in academic finance and is more '
        'accurate for very low interest rates.')
    add_body_text(doc, 'Formula: Years to Double = 69.3 ÷ Interest Rate (%)')

    add_styled_heading(doc, 'The Rule of 70', level=2)
    add_body_text(doc,
        'The Rule of 70 is a compromise between the mathematical precision of 69.3 and the '
        'mental-math convenience of 72. It is commonly used by economists when discussing GDP growth '
        'rates, population growth, and inflation.')
    add_body_text(doc, 'Formula: Years to Double = 70 ÷ Interest Rate (%)')

    add_styled_heading(doc, 'Which Rule Should You Use?', level=2)

    add_branded_table(doc,
        headers=['Rule', 'Best For', 'Accuracy'],
        data=[
            ('Rule of 69.3', 'Continuous compounding,\nacademic calculations', 'Most precise mathematically'),
            ('Rule of 70', 'Low rates (1-5%),\neconomics/inflation', 'Good all-round accuracy'),
            ('Rule of 72', 'General investing (6-10%),\nquick mental maths', 'Best ease-of-use,\nmost divisors'),
        ],
        header_color='DB2777'
    )

    add_body_text(doc,
        'For most everyday investors, the Rule of 72 remains the best choice due to its '
        'simplicity and the fact that 72 has so many divisors.')
    doc.add_page_break()

    # ===== CHAPTER 9: LIMITATIONS =====
    add_styled_heading(doc, '9. Limitations and When Not to Use It', level=1)

    add_body_text(doc, 'While the Rule of 72 is fantastic, it is important to understand its limitations:')

    add_styled_heading(doc, 'Variable Interest Rates', level=2)
    add_body_text(doc,
        'The Rule assumes a constant rate of return. In reality, investment returns fluctuate. '
        'Stock markets might return 20% one year and -10% the next. The rule works best with '
        'average expected returns over long periods.')

    add_styled_heading(doc, 'Very High or Very Low Rates', level=2)
    add_body_text(doc,
        'The rule loses accuracy at extreme rates. Below 2%, use the Rule of 70 or 69.3. '
        'Above 20%, the approximation breaks down significantly.')

    add_styled_heading(doc, 'Taxes and Fees', level=2)
    add_body_text(doc,
        'The Rule does not account for taxes, management fees, or transaction costs. '
        'If your investment returns 8% but you pay 2% in fees and taxes, your effective rate is '
        '6%. Always use the net (after-fee, after-tax) rate for realistic estimates.')

    add_styled_heading(doc, 'Inflation', level=2)
    add_body_text(doc,
        'The rule calculates nominal doubling, not real purchasing power. '
        'To estimate real doubling time, subtract the inflation rate from your return. '
        'Example: 8% return - 3% inflation = 5% real return. 72 ÷ 5 = 14.4 years to double in real terms.')

    add_styled_heading(doc, 'Not a Guarantee', level=2)
    add_body_text(doc,
        'The Rule of 72 is an estimation tool, not a promise. Past returns do not guarantee future '
        'results. Always use it as a planning guide, not an investment guarantee.')

    add_highlight_box(doc, 'Always use your NET return (after fees, taxes, and inflation)\nfor the most realistic Rule of 72 estimate.')
    doc.add_page_break()

    # ===== CHAPTER 10: CONCLUSION =====
    add_styled_heading(doc, '10. Key Takeaways and Conclusion', level=1)

    add_body_text(doc,
        'The Rule of 72 is one of the simplest yet most powerful concepts in personal finance. '
        'Here are the key points to remember:')

    for t in [
        'Divide 72 by the annual rate of return to estimate how long until your investment doubles.',
        'It works best for interest rates between 2% and 20%.',
        'Even a 1-2% difference in annual returns has a massive impact over decades.',
        'Use it in reverse to understand how quickly inflation erodes purchasing power.',
        'Compound interest is exponential — each doubling adds more absolute wealth than all previous doublings combined.',
        'Use your net return after subtracting fees, taxes, and inflation for the most accurate estimates.',
        'Start investing as early as possible — time is the most important ingredient in compounding.',
        'The Rule of 72 is a planning tool, not a guarantee. Always seek professional financial advice.',
    ]:
        para = doc.add_paragraph(t, style='List Bullet')
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    add_styled_heading(doc, 'Your Next Steps', level=2)

    add_body_text(doc,
        'Now that you understand the Rule of 72, put it to work! Use it to evaluate investment '
        'opportunities, compare savings accounts, plan for retirement, or appreciate the power '
        'of compound growth. The sooner you start, the more doublings your money will achieve.')

    add_body_text(doc,
        'Visit Income Online at www.incomeonline.info to discover 199+ legitimate platforms '
        'where you can start growing your income today. From freelancing to investing, from '
        'e-commerce to teaching — there is an earning opportunity that suits your skills and goals.')

    add_closing_page(doc)

    return save_to_buffer(doc)


if __name__ == '__main__':
    buf = generate_rule72_document()
    with open('/app/backend/static/The_Rule_of_72_Guide.docx', 'wb') as f:
        f.write(buf.read())
    print('Document generated successfully!')
