"""Generate the Beginner's Guide to Passive Income using the MoneyRules template."""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, DEEP_PURPLE, PURPLE, PINK, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_passive_income_document():
    doc = create_moneyrules_document(
        title="Beginner's Guide to Passive Income",
        subtitle='How Ordinary People Build Money-Making Assets'
    )

    add_title_page(doc,
        title="Passive Income",
        subtitle="A Beginner's Guide to Earning Money While You Sleep",
        tagline='Seven proven income streams, what they really earn,\nand how to start each one with whatever you have today.')

    # TOC
    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()
    for num, t in [
        ('1.', 'Introduction — What Passive Income Really Means'),
        ('2.', 'The Great Myth: "Passive" Does Not Mean "Free"'),
        ('3.', 'The Seven Streams of Passive Income'),
        ('4.', 'Stream 1: Dividend Stocks & Index Funds'),
        ('5.', 'Stream 2: High-Yield Savings & Bonds'),
        ('6.', 'Stream 3: Rental Property & REITs'),
        ('7.', 'Stream 4: Digital Products & Royalties'),
        ('8.', 'Stream 5: Affiliate & Content Revenue'),
        ('9.', 'Starter Strategies by Budget'),
        ('10.', 'Common Mistakes & Your Next Steps'),
    ]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(10)
        r = para.add_run(f'{num}  '); r.font.size = Pt(13); r.font.color.rgb = PURPLE; r.bold = True
        r2 = para.add_run(t); r2.font.size = Pt(13); r2.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run('DISCLAIMER: This guide is for educational purposes only and does not constitute financial advice. Always consult a qualified adviser before investing.')
    run.font.size = Pt(9); run.font.color.rgb = GREY; run.italic = True
    doc.add_page_break()

    # Ch 1
    add_styled_heading(doc, '1. Introduction — What Passive Income Really Means', level=1)
    add_body_text(doc, 'Passive income is money that keeps arriving even when you\'re not actively working for it. Your job is "active" income: no work, no pay. Passive income is different — you do the work once (or invest the money once), and the returns flow in month after month, year after year.')
    add_body_text(doc, 'Classic examples include rent from a property you own, dividends from shares you hold, royalties from a book you wrote, or advertising revenue from a blog post you published three years ago. In every case, the work-to-earning link is broken.')
    add_highlight_box(doc, 'Active income pays for your time.\nPassive income pays for your past decisions.')
    add_styled_heading(doc, 'Why Does It Matter?', level=2)
    add_body_text(doc, 'The fundamental limitation of active income is that you only have 24 hours in a day. If you earn £30/hour, you can never out-earn £30 × 24 × 365 per year, no matter how hard you work. Passive income has no such ceiling. It scales with capital, audience, or intellectual property — not with your time.')
    add_body_text(doc, 'The goal of this guide is not to promise that you can quit your job next month and live off passive income. It is to show you how ordinary people start one, two, or three streams with what they have today — and let those streams grow in the background over years.')
    doc.add_page_break()

    # Ch 2
    add_styled_heading(doc, '2. The Great Myth: "Passive" Does Not Mean "Free"', level=1)
    add_body_text(doc, 'Internet gurus love to sell passive income as if it were a magic escalator: step on, do nothing, arrive at wealth. It is not. Every legitimate passive income stream requires one of two things upfront — money (capital) or effort (time and expertise). Usually both.')
    add_styled_heading(doc, 'The Honest Framework', level=2)
    add_branded_table(doc,
        headers=['Requires', 'Examples', 'Typical Setup Time'],
        data=[
            ('Money upfront', 'Dividend stocks, bonds, rental property, REITs', 'Days'),
            ('Effort upfront', 'Writing a book, building a website, creating a course', '3 months – 2 years'),
            ('Both', 'A rental property you also renovate, a self-published best-seller with paid ads', '6 months – 5 years'),
        ])
    add_body_text(doc, 'Once the asset is built or bought, the ongoing work is minimal — maybe a few hours a month of maintenance. But "minimal" is not zero. Be suspicious of anyone who claims otherwise.')
    add_highlight_box(doc, 'All passive income is front-loaded with work or money.\nThe "passive" part is what comes afterwards.')
    doc.add_page_break()

    # Ch 3
    add_styled_heading(doc, '3. The Seven Streams of Passive Income', level=1)
    add_body_text(doc, 'Most of what is sold as "passive income" falls into one of seven categories. We\'ll explore each in turn.')
    add_branded_table(doc,
        headers=['Stream', 'Capital Needed', 'Time to Profit', 'Risk Level'],
        data=[
            ('1. Dividend stocks / ETFs',   '£100+',          'Immediate',  'Medium'),
            ('2. High-yield savings / bonds','£1+',            'Immediate',  'Low'),
            ('3. Rental property',           '£30,000+',       '1–3 months',  'High'),
            ('4. REITs (Real Estate Funds)', '£100+',          'Immediate',  'Medium'),
            ('5. Digital products',          '£0–£500',        '3–12 months', 'Medium'),
            ('6. Affiliate / blog revenue',  '£0–£200',        '6–24 months', 'Medium'),
            ('7. Peer-to-peer lending',      '£100+',          '1 month',    'High'),
        ],
        header_color='DB2777')
    add_body_text(doc, 'Most smart investors build a portfolio of 2–4 of these — never all seven. Diversification across streams is more important than picking the "best" one.')
    doc.add_page_break()

    # Ch 4
    add_styled_heading(doc, '4. Stream 1: Dividend Stocks & Index Funds', level=1)
    add_body_text(doc, 'When you buy shares in a company, you become a part-owner. Profitable companies often pay out a portion of their profits to shareholders every quarter — these are dividends. Over the long run, UK and US dividend index funds have yielded around 3–5% per year in cash payments, plus capital appreciation of the share price itself.')
    add_styled_heading(doc, 'How to Start (5 Steps)', level=2)
    for t in [
        'Open a Stocks & Shares ISA with a low-cost platform (e.g. Vanguard, InvestEngine, Trading 212).',
        'Pick a diversified index fund — FTSE Global All Cap or S&P 500 are popular starters.',
        'Set up a monthly direct debit (even £50/month compounds).',
        'Turn ON automatic dividend reinvestment — "Accumulation" share class.',
        'Leave it alone for 10+ years. Seriously — do not trade.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_highlight_box(doc, 'A £200/month ISA contribution at 7% average return\nbecomes roughly £104,000 after 20 years.')
    doc.add_page_break()

    # Ch 5
    add_styled_heading(doc, '5. Stream 2: High-Yield Savings & Bonds', level=1)
    add_body_text(doc, 'The simplest passive income of all: put money in a savings account that pays interest. In 2026, top UK easy-access savings accounts pay around 4–5% AER, and premium bonds offer tax-free prize draws.')
    add_styled_heading(doc, 'The Options, Ranked by Effort', level=2)
    add_branded_table(doc,
        headers=['Option', '2026 Yield', 'Access'],
        data=[
            ('Easy-access savings',  '3.5 – 5.0%', 'Instant'),
            ('1-year fixed savings', '4.5 – 5.5%', 'Locked 12 months'),
            ('Cash ISA',             '4.0 – 5.0%', 'Instant (tax-free)'),
            ('UK Gilts / Bonds',     '4.0 – 5.5%', '2–30 year maturity'),
            ('Premium Bonds',        '~4.0% avg',  'Instant'),
        ])
    add_body_text(doc, 'This is your emergency fund territory. The returns are modest, but the capital is virtually guaranteed — and that safety is exactly what a diversified portfolio needs.')
    doc.add_page_break()

    # Ch 6
    add_styled_heading(doc, '6. Stream 3: Rental Property & REITs', level=1)
    add_body_text(doc, 'Owning property you let out is the most traditional passive income stream. A typical buy-to-let in the UK yields around 4–7% gross, with the property itself historically appreciating 2–4% per year on top.')
    add_body_text(doc, 'The downside: it is NOT very passive. Tenants break things. Boilers fail at 2am. Tax rules for landlords have tightened significantly since 2020.')
    add_styled_heading(doc, 'REITs: The Passive Alternative', level=2)
    add_body_text(doc, 'A Real Estate Investment Trust (REIT) is effectively a stock-market fund that owns property. You buy shares like any other investment, receive dividends quarterly from the rental income, and never deal with a single tenant. UK REITs such as Tritax Big Box or British Land yield 4–6% annually with zero maintenance responsibility.')
    add_highlight_box(doc, 'REITs give you the income of a landlord\nwithout the 2am calls from a burst pipe.')
    doc.add_page_break()

    # Ch 7
    add_styled_heading(doc, '7. Stream 4: Digital Products & Royalties', level=1)
    add_body_text(doc, 'This is the "build it once, sell it forever" stream. Write a book, create an online course, design printable templates, compose a piece of stock music. Each new customer costs you almost nothing to serve.')
    add_styled_heading(doc, 'Realistic Earning Ranges', level=2)
    add_branded_table(doc,
        headers=['Product Type', 'Effort', 'Typical Monthly Income After 1 Year'],
        data=[
            ('Self-published Kindle book', '200–400 hrs writing', '£50 – £1,500'),
            ('Online course (Udemy/Teachable)', '100–200 hrs', '£100 – £3,000'),
            ('Printable templates (Etsy)', '20–100 hrs', '£20 – £500'),
            ('Stock photos / music', '50–300 hrs', '£30 – £400'),
            ('Mobile app with ads', '200–500 hrs', '£0 – £2,000'),
        ],
        header_color='EA580C')
    add_body_text(doc, 'The distribution is highly skewed: 80% of products earn almost nothing, and 20% earn almost everything. Plan for the 80% case. Treat the 20% case as a bonus.')
    doc.add_page_break()

    # Ch 8
    add_styled_heading(doc, '8. Stream 5: Affiliate & Content Revenue', level=1)
    add_body_text(doc, 'Write helpful blog content, YouTube videos, or TikTok tutorials. Recommend products you genuinely like. When viewers click your tracked link and buy, you earn a commission — typically 3–20% of the sale.')
    add_body_text(doc, 'This is arguably the hardest stream. It is dependent on search algorithms and platform whims, and it takes 12–24 months to build meaningful traffic. But once established, a well-optimised content site can produce £500–£5,000/month with almost no ongoing work besides occasional refreshes.')
    add_highlight_box(doc, 'Content businesses are brutal at the start\nand beautiful at the end.')
    doc.add_page_break()

    # Ch 9
    add_styled_heading(doc, '9. Starter Strategies by Budget', level=1)
    add_styled_heading(doc, 'Budget: £500 or less', level=2)
    add_body_text(doc, 'Open a Stocks & Shares ISA, buy a global index fund accumulation class, set up £50–100/month. Forget about it. This will not make you rich this year, but it will silently compound for the next 30.')
    add_styled_heading(doc, 'Budget: £500 – £5,000', level=2)
    add_body_text(doc, 'Max out a cash ISA for a 3-month emergency fund, then start a Stocks & Shares ISA with the rest. Add a small position in a REIT (5–10% of portfolio) for property exposure.')
    add_styled_heading(doc, 'Budget: £5,000 – £50,000', level=2)
    add_body_text(doc, 'Diversify across dividend ETFs, bonds, REITs, and premium bonds. Consider starting a digital side-product (course, Kindle book) while your investments compound.')
    add_styled_heading(doc, 'Budget: £50,000+', level=2)
    add_body_text(doc, 'Now buy-to-let becomes realistic, as does a dedicated portfolio of 4–6 dividend stocks picked individually. Consider a SIPP for tax-advantaged retirement investing.')
    doc.add_page_break()

    # Ch 10
    add_styled_heading(doc, '10. Common Mistakes & Your Next Steps', level=1)
    add_body_text(doc, 'Five mistakes that kill most passive income attempts:')
    for t in [
        'Jumping between strategies — give each at least 18 months before switching.',
        'Starting too big — failing on a £200/month ISA is survivable; failing on a £40,000 buy-to-let is not.',
        'Forgetting taxes — dividends above £500/year are taxable in 2026; factor this into net returns.',
        'Chasing high yields — anything promising 15%+ guaranteed is almost certainly a scam.',
        'Stopping the contributions — compounding only works when you keep feeding it.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT

    add_styled_heading(doc, 'Your First Three Steps This Week', level=2)
    for t in [
        'Day 1: Open a Stocks & Shares ISA (takes ~10 minutes online).',
        'Day 2: Set up a £50/month direct debit into a global index fund.',
        'Day 3: Pick ONE content or product idea and block 2 hours a week to build it.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT

    add_body_text(doc, 'Visit www.incomeonline.info to discover 199+ legitimate platforms that can accelerate every one of these streams.')
    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    buf = generate_passive_income_document()
    os.makedirs('/app/backend/static', exist_ok=True)
    with open('/app/backend/static/Passive_Income_Beginners_Guide.docx', 'wb') as f:
        f.write(buf.read())
    print('Passive Income guide generated!')
