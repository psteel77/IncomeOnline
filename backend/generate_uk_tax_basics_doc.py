"""Generate the UK Tax Basics for Freelancers Guide using the MoneyRules template."""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, PURPLE, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_uk_tax_basics_document():
    doc = create_moneyrules_document(
        title='UK Tax Basics for Freelancers & Side-Hustlers',
        subtitle='Self Assessment Demystified — Without the Accountant Jargon'
    )
    add_title_page(doc,
        title='UK Tax Basics',
        subtitle='For Freelancers, Side-Hustlers and the Self-Employed',
        tagline='When to register, what you can claim, key dates,\nand how to avoid painful tax-year surprises.')

    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()
    for num, t in [
        ('1.', 'Introduction — Why Tax Matters Before You Think It Does'),
        ('2.', 'The £1,000 Trading Allowance'),
        ('3.', 'When Do I Need to Register for Self Assessment?'),
        ('4.', 'National Insurance Contributions Explained'),
        ('5.', 'Allowable Expenses You Can Claim'),
        ('6.', 'Record-Keeping: What HMRC Actually Wants'),
        ('7.', 'Key Dates and Deadlines'),
        ('8.', 'Payments on Account — The January Surprise'),
        ('9.', 'When to Hire an Accountant'),
        ('10.', 'Key Takeaways and Your Action Plan'),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f'{num}  '); r.font.size = Pt(13); r.font.color.rgb = PURPLE; r.bold = True
        r2 = p.add_run(t); r2.font.size = Pt(13); r2.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run('DISCLAIMER: General educational content only, reflecting 2026 UK rules. Not personal tax advice. Always consult HMRC directly or a qualified accountant.')
    run.font.size = Pt(9); run.font.color.rgb = GREY; run.italic = True
    doc.add_page_break()

    add_styled_heading(doc, '1. Introduction — Why Tax Matters Before You Think It Does', level=1)
    add_body_text(doc, 'If you\'re earning anything outside PAYE — freelance gigs, Etsy sales, a YouTube channel, survey income, crypto, eBay reselling — you are potentially trading as a self-employed person in the eyes of HMRC. The threshold is lower than most people realise, and the penalties for ignoring it are harsh.')
    add_body_text(doc, 'The good news: the UK tax system for small earners is actually relatively friendly. You get a tax-free trading allowance, generous expense deductions, and simple online filing. The bad news: HMRC is increasingly using data-matching (eBay, Airbnb, PayPal, Upwork all now report to HMRC) to find people who aren\'t declaring.')
    add_highlight_box(doc, 'Register on time → pay normal tax.\nGet caught unregistered → tax + penalties up to 100% of the tax owed.')
    doc.add_page_break()

    add_styled_heading(doc, '2. The £1,000 Trading Allowance', level=1)
    add_body_text(doc, 'Since 2017, every UK individual has a £1,000 "trading allowance" — tax-free income from self-employment each tax year (6 April to 5 April). You do not need to register for Self Assessment if your TOTAL self-employed gross income for the year is £1,000 or less.')
    add_styled_heading(doc, 'Worked examples', level=2)
    add_branded_table(doc,
        headers=['Your side income this year', 'Action Required?'],
        data=[
            ('£600 from Fiverr gigs',                            'NO — below £1,000'),
            ('£950 from YouTube AdSense',                        'NO — below £1,000 (just)'),
            ('£1,100 from eBay reselling',                       'YES — register'),
            ('£500 Fiverr + £600 Etsy = £1,100 total',          'YES — it\'s the TOTAL that counts'),
            ('£4,000 from Upwork as a full freelance gig',      'YES — register'),
        ])
    add_body_text(doc, 'Note: even if you are under £1,000, you can OPTIONALLY register (e.g. to build a State Pension record or claim losses). But you don\'t have to.')
    doc.add_page_break()

    add_styled_heading(doc, '3. When Do I Need to Register for Self Assessment?', level=1)
    add_body_text(doc, 'You must register with HMRC by 5 October after the end of the tax year in which you exceeded £1,000. Example: you earned £1,500 via freelancing between April 2025 and April 2026. Deadline to register: 5 October 2026. Deadline to file your first return: 31 January 2027.')
    add_styled_heading(doc, 'How to Register', level=2)
    for t in [
        'Go to gov.uk/register-for-self-assessment.',
        'Fill out form CWF1 (self-employed).',
        'You\'ll receive a 10-digit Unique Taxpayer Reference (UTR) by post — guard this, you need it to file.',
        'Also get a Gateway ID & password for the HMRC online service.',
        'Process takes ~2–4 weeks. Don\'t leave it to the last minute.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_highlight_box(doc, 'Register EARLY. Waiting until late January\nis the #1 cause of missed-deadline penalties.')
    doc.add_page_break()

    add_styled_heading(doc, '4. National Insurance Contributions Explained', level=1)
    add_body_text(doc, 'As a self-employed person, you pay two types of NI in addition to Income Tax:')
    add_branded_table(doc,
        headers=['Type', '2026 Rate', 'When It Applies'],
        data=[
            ('Class 2 NI',  'Voluntary',  'Abolished in 2024 for most. Voluntary if profits < £6,725 to protect State Pension record.'),
            ('Class 4 NI',  '6% / 2%',    '6% on profits £12,570–£50,270. 2% on anything above £50,270.'),
        ])
    add_body_text(doc, 'NI is calculated on your PROFIT (income minus expenses), not on gross income. It\'s collected automatically as part of your Self Assessment — no separate filing.')
    add_highlight_box(doc, 'Under £6,725 profit? Consider paying voluntary Class 2 (£3.45/wk)\nto keep building State Pension credits.')
    doc.add_page_break()

    add_styled_heading(doc, '5. Allowable Expenses You Can Claim', level=1)
    add_body_text(doc, 'You pay tax on PROFIT, not on revenue. Every legitimate business expense you deduct directly reduces your tax bill. For a basic-rate taxpayer, every £1 of claimed expense saves 29p (20% tax + 6% Class 4 + roughly 3% overhead).')
    add_branded_table(doc,
        headers=['Expense Category', 'Examples'],
        data=[
            ('Home office',      'Flat rate £6/wk OR % of bills if >25hrs/month from home'),
            ('Technology',       'Laptop, phone (business %), software subscriptions'),
            ('Travel',           'Mileage @45p/mi for first 10k, train/bus fares'),
            ('Professional',     'Accountant fees, business bank charges, courses'),
            ('Marketing',        'Website hosting, domain, ads, business cards'),
            ('Subscriptions',    'Trade magazines, membership fees, platform fees'),
            ('Stationery',       'Printer paper, ink, postage'),
        ],
        header_color='DB2777')
    add_body_text(doc, 'Golden rule: the expense must be "wholly and exclusively" for business. A laptop used 70% for work, 30% personal = claim 70%. A Netflix subscription you use "sometimes for research" = DON\'T claim.')
    doc.add_page_break()

    add_styled_heading(doc, '6. Record-Keeping: What HMRC Actually Wants', level=1)
    add_body_text(doc, 'HMRC requires you to keep records for at least 5 years after the filing deadline. For simple freelancers, this can be a spreadsheet plus a folder of receipts. You do NOT need commercial accounting software unless your turnover exceeds £90,000 (VAT threshold) or you file under Making Tax Digital.')
    add_styled_heading(doc, 'Minimum Records to Keep', level=2)
    for t in [
        'All sales invoices or self-billing records (what you charged, when, to whom).',
        'All purchase receipts or online order confirmations.',
        'Business bank/PayPal/Stripe statements.',
        'Mileage log for business car journeys.',
        'Home office hours log if using the actual-cost method.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_highlight_box(doc, 'Open a SEPARATE bank account for your self-employment.\nIt makes records trivial, and it\'s free at most UK banks.')
    doc.add_page_break()

    add_styled_heading(doc, '7. Key Dates and Deadlines', level=1)
    add_branded_table(doc,
        headers=['Date', 'Event', 'Miss It → Penalty'],
        data=[
            ('6 April',        'New tax year starts',                               '—'),
            ('5 April next',   'Tax year ends',                                     '—'),
            ('5 October',      'Deadline to register if first year of trading',     '£100 + %age'),
            ('31 October',     'Paper return deadline (rarely used now)',           '£100'),
            ('31 January',     'Online return deadline + final tax payment',        '£100 + daily £10'),
            ('31 July',        'Payment on account (2nd instalment)',               'Interest on arrears'),
        ],
        header_color='EA580C')
    add_body_text(doc, 'The 31 January deadline is the BIG one. Missing it by a day = instant £100 penalty, regardless of whether you owe anything. Miss it by 3 months and you accrue £10/day fines, capped at £900 plus percentages of tax owed.')
    doc.add_page_break()

    add_styled_heading(doc, '8. Payments on Account — The January Surprise', level=1)
    add_body_text(doc, 'This catches almost every new freelancer. If your tax bill in year 1 exceeds £1,000, HMRC asks you to pay TWO things on 31 January: the tax you owe PLUS half of next year\'s estimated tax.')
    add_styled_heading(doc, 'Worked Example', level=2)
    add_body_text(doc, 'Your 2025/26 tax return (filed by 31 Jan 2027) shows a tax bill of £3,000. On 31 January 2027 you owe:')
    for t in [
        '£3,000 — the actual tax for the year just ended.',
        '£1,500 — half of 2026/27 tax, estimated at the same £3,000.',
        'TOTAL DUE 31 JAN: £4,500.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_body_text(doc, 'Then on 31 July 2027 you owe the OTHER £1,500. In effect, you\'re paying 150% of a year\'s tax in January of year 2. This is the "January surprise" that sinks many freelancers\' finances.')
    add_highlight_box(doc, 'From day 1 of freelancing, save 30% of every payment\ninto a "tax pot" account. You\'ll thank yourself in January.')
    doc.add_page_break()

    add_styled_heading(doc, '9. When to Hire an Accountant', level=1)
    add_body_text(doc, 'For most side-hustlers earning £1,000–£30,000/year, DIY via HMRC\'s online Self Assessment is perfectly adequate. It\'s free, the form is straightforward, and the interface is surprisingly well-designed.')
    add_styled_heading(doc, 'Hire an accountant if you:', level=2)
    for t in [
        'Earn over £50,000+ profit and want tax-efficient planning.',
        'Are considering forming a limited company.',
        'Have complex income (multiple sources, international clients, crypto, rental).',
        'Have £5,000+ of equipment or capital allowances.',
        'Are subject to an HMRC enquiry — get professional help immediately.',
        'Simply value your time at more than £30/hour (a simple return takes ~6 hours).',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_body_text(doc, 'Typical fees in 2026: £200–£400 for a simple Self Assessment, £600–£1,200 annually for full ltd company accounts. In most cases the accountant saves you more than their fee by spotting allowances you\'d miss.')
    doc.add_page_break()

    add_styled_heading(doc, '10. Key Takeaways and Your Action Plan', level=1)
    for t in [
        '£1,000 trading allowance — below this, no registration needed.',
        'Register for Self Assessment by 5 October after you cross the threshold.',
        'Tax is on PROFIT not revenue — track every allowable expense.',
        'Separate bank account + spreadsheet = adequate records for most freelancers.',
        '31 January is the deadline for everything. Missing it costs £100 minimum.',
        'Save 30% of every freelance payment into a tax pot.',
        'Payments on account = you may owe 150% in January of year 2. Plan for this.',
        'Hire an accountant once your income or complexity justifies £300+/year.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT

    add_styled_heading(doc, 'Your First Two Weeks', level=2)
    for t in [
        'Day 1: Add up last 12 months of side income. Are you above £1,000?',
        'Day 2–3: If yes, register for Self Assessment at gov.uk/register-for-self-assessment.',
        'Day 4: Open a separate "business" current account (Starling or Monzo business are popular).',
        'Day 5: Start a simple spreadsheet — Date / Description / Income / Expense / Category.',
        'Day 6: Open a separate savings account labelled "TAX". Set up a 30% rule.',
        'Day 7–14: Back-fill your records for the current tax year to date.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_body_text(doc, 'Visit www.incomeonline.info to discover 199+ legitimate platforms for freelancing, gigs and remote work — each one potentially a tax-deductible business activity.')
    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    buf = generate_uk_tax_basics_document()
    os.makedirs('/app/backend/static', exist_ok=True)
    with open('/app/backend/static/UK_Tax_Basics_Freelancers.docx', 'wb') as f:
        f.write(buf.read())
    print('UK Tax Basics guide generated!')
