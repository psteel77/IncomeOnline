"""
Blank branded MoneyRules template — the IncomeOnline house style with NO content.

Produces a skeleton the client can fill in: title page, branded chapter panels
(each starting on a new page), one of every coloured callout with placeholder
text, and a closing page. Page borders, running header, footer, Montserrat bold
throughout — identical formatting to the finished guides.

Outputs both a .docx (editable master) and a .pdf are built from this single
source, so the Word and PDF templates always match.
"""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_expert_tip, add_action_checklist, add_common_mistake,
    add_example_box, add_case_study, add_closing_page, save_to_buffer, BODY_TEXT, GREY,
)
from docx.shared import Pt


def _placeholder(doc, text):
    """A faint italic instruction line the author can delete/replace."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.italic = True
    r.bold = False
    r.font.color.rgb = GREY
    return p


def _blank_bullets(doc, n=3):
    for _ in range(n):
        p = doc.add_paragraph('[ List item ]', style='List Bullet')
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.color.rgb = BODY_TEXT


def generate_template_document():
    doc = create_moneyrules_document(
        title='[ Your Guide Title ]',
        subtitle='PILLAR [ # ]',
    )
    add_title_page(
        doc,
        title='[ Your Guide Title ]',
        subtitle='PILLAR [ # ]',
        tagline='[ One line describing what this guide helps the reader do ]',
    )

    # ---- Chapter 1: shows every building block available ----
    add_styled_heading(doc, 'Chapter 1 – [ Chapter Title ]', level=1)
    _placeholder(doc, 'Replace the placeholder text below with your own content. '
                      'Each chapter automatically starts on a new page with a branded heading.')
    add_body_text(doc, '[ Write your opening paragraph here. Keep paragraphs short and clear — '
                       'this text is Montserrat 11pt, bold, in the IncomeOnline house style. ]')
    add_body_text(doc, '[ A second paragraph. Add as many as you need. ]')

    add_styled_heading(doc, '[ A Section Heading ]', level=2)
    add_body_text(doc, '[ Body text for this section. ]')
    _blank_bullets(doc, 3)

    add_expert_tip(doc, '[ Expert Tip: a short, high-value pointer for the reader. This box, and every '
                        'coloured box, stays together and moves to the next page if it will not fit. ]')
    add_example_box(doc, '[ Example: a worked example or short scenario that illustrates the point. ]')
    add_action_checklist(doc, [
        '[ Action one the reader should take ]',
        '[ Action two ]',
        '[ Action three ]',
    ])
    add_common_mistake(doc, '[ Common Mistake: the pitfall to avoid, and what to do instead. ]')
    add_case_study(doc, '[ IncomeOnline Case Study: a real-world illustration that builds trust. ]')

    # ---- Chapters 2–4: blank chapter pages ready to fill ----
    for n in (2, 3, 4):
        add_styled_heading(doc, f'Chapter {n} – [ Chapter Title ]', level=1)
        _placeholder(doc, '[ This chapter starts on a fresh page. Add your text, headings and '
                          'coloured boxes here. ]')
        add_body_text(doc, '[ … ]')

    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    buf = generate_template_document()
    with open('/app/backend/static/MoneyRules_Branded_Template.docx', 'wb') as f:
        f.write(buf.read())
    print('Branded template generated successfully!')
