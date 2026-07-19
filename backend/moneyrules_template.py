"""
MoneyRules — professional branded Word document template for Income Online.

Print-brochure quality, NOT a plain Word document. Applies the IncomeOnline
house rules requested by the client:

  1. Coloured panels never sit split/orphaned at the bottom of a page — every
     callout is a single-row table with `w:cantSplit`, so the whole block moves
     to the top of the next page if it does not fit.
  2. Expert Tip / Action Checklist / Common Mistake / Example / Case Study blocks
     stay together with their heading (kept as one unbreakable unit).
  3. Every chapter (Heading 1) renders as a full-width branded purple panel and
     chapters start on a new page (page breaks handled by the generators).
  4. IncomeOnline colour scheme throughout — purple / pink / orange.
  5. Montserrat 11pt body text throughout.
  6. Page borders on every page.
  7. Branded chapter headings + running "INCOME ONLINE" header + page-number footer.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from io import BytesIO


# ---------------------------------------------------------------
# Brand palette — purple / pink / orange
# ---------------------------------------------------------------
DEEP_PURPLE = RGBColor(0x4C, 0x1D, 0x95)     # purple-900
PURPLE = RGBColor(0x6D, 0x28, 0xD9)           # purple-700
MEDIUM_PURPLE = RGBColor(0x7C, 0x3A, 0xED)    # purple-600
PINK = RGBColor(0xDB, 0x27, 0x77)             # pink-600
ROSE = RGBColor(0xEC, 0x48, 0x99)             # pink-500
ORANGE = RGBColor(0xEA, 0x58, 0x0C)           # orange-600
DARK_TEXT = RGBColor(0x1E, 0x1B, 0x4B)         # indigo-950
BODY_TEXT = RGBColor(0x14, 0x14, 0x18)         # near-black (bold, high-contrast)
GREY = RGBColor(0x64, 0x74, 0x8B)             # slate-500
LIGHT_GREY = RGBColor(0x94, 0xA3, 0xB8)       # slate-400
ACCENT_GOLD = RGBColor(0xB4, 0x5D, 0x09)      # amber-700 (kept for back-compat)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = 'Montserrat'

# Hex helpers for shading / borders
PURPLE_HEX = '6D28D9'
DEEP_PURPLE_HEX = '4C1D95'
PINK_HEX = 'DB2777'
ORANGE_HEX = 'EA580C'
TINT_PURPLE = 'F3F0FF'   # light purple wash
TINT_PINK = 'FDF2F8'     # light pink wash
TINT_ORANGE = 'FFF4ED'   # light orange wash

BORDER_OUTER = DEEP_PURPLE_HEX
BORDER_INNER = PINK_HEX


# ---------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------

def _apply_font(font, name=FONT):
    """Force a run/style font family across ascii/hAnsi/cs so LibreOffice honours it."""
    font.name = name
    rpr = font.element.get_or_add_rPr() if hasattr(font.element, 'get_or_add_rPr') else None


def _set_run_font(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), name)


def _set_style_font(style, name=FONT):
    try:
        style.font.name = name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rfonts.set(qn(attr), name)
    except Exception:
        pass


def _cant_split(row):
    """Stop a table row splitting across a page → single-row callouts move wholesale."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))


def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex, qn('w:val'): 'clear',
    })
    tcPr.append(shd)


def _cell_margins(cell, top=140, bottom=140, start=180, end=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = tcPr.makeelement(qn('w:tcMar'), {})
    for side, val in (('top', top), ('bottom', bottom), ('start', start), ('end', end)):
        tcMar.append(tcMar.makeelement(qn(f'w:{side}'), {qn('w:w'): str(val), qn('w:type'): 'dxa'}))
    tcPr.append(tcMar)


def _set_cell_borders(cell, color='4C1D95', size='4'):
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right'):
        borders.append(borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): size, qn('w:color'): color, qn('w:space'): '0',
        }))
    tcPr.append(borders)


def _accent_left_border(cell, accent_hex, tint_hex):
    """Thick coloured left rule + hairline tinted box — the callout look."""
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    borders.append(borders.makeelement(qn('w:left'), {
        qn('w:val'): 'single', qn('w:sz'): '42', qn('w:color'): accent_hex, qn('w:space'): '0',
    }))
    for edge in ('top', 'bottom', 'right'):
        borders.append(borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:color'): accent_hex, qn('w:space'): '0',
        }))
    tcPr.append(borders)


def _full_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.makeelement(qn('w:tblW'), {qn('w:w'): '5000', qn('w:type'): 'pct'})
    tblPr.append(tblW)


# ---------------------------------------------------------------
# Page furniture: borders, running header, footer
# ---------------------------------------------------------------

def _add_page_borders(section):
    sectPr = section._sectPr
    pgBorders = sectPr.makeelement(qn('w:pgBorders'), {qn('w:offsetFrom'): 'page'})
    cfg = {
        'top':    {'val': 'thinThickSmallGap', 'sz': '30', 'space': '20', 'color': BORDER_OUTER},
        'bottom': {'val': 'thickThinSmallGap', 'sz': '30', 'space': '20', 'color': BORDER_OUTER},
        'left':   {'val': 'thinThickSmallGap', 'sz': '30', 'space': '20', 'color': BORDER_OUTER},
        'right':  {'val': 'thickThinSmallGap', 'sz': '30', 'space': '20', 'color': BORDER_OUTER},
    }
    for edge, c in cfg.items():
        pgBorders.append(pgBorders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): c['val'], qn('w:sz'): c['sz'], qn('w:space'): c['space'], qn('w:color'): c['color'],
        }))
    sectPr.append(pgBorders)


def _add_running_header(section):
    """Small centred 'INCOME ONLINE' brand line at the top of every page (not the cover)."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('I N C O M E   O N L I N E')
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = MEDIUM_PURPLE
    _set_run_font(run)
    # thin orange rule under the header
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    pBdr.append(pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '3', qn('w:color'): ORANGE_HEX,
    }))
    pPr.append(pBdr)


def _add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    rule_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    rule_para.clear()
    pPr = rule_para._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    pBdr.append(pBdr.makeelement(qn('w:top'), {
        qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '4', qn('w:color'): ORANGE_HEX,
    }))
    pPr.append(pBdr)
    rule_para.paragraph_format.space_before = Pt(0)
    rule_para.paragraph_format.space_after = Pt(2)

    tabs = pPr.makeelement(qn('w:tabs'), {})
    tabs.append(tabs.makeelement(qn('w:tab'), {qn('w:val'): 'center', qn('w:pos'): '4680'}))
    tabs.append(tabs.makeelement(qn('w:tab'), {qn('w:val'): 'right', qn('w:pos'): '9360'}))
    pPr.append(tabs)

    def _pgfont(r):
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        _set_run_font(r)

    run_pg_label = rule_para.add_run('Page ')
    _pgfont(run_pg_label)
    fld1 = run_pg_label._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_pg = rule_para.add_run(); run_pg._element.append(fld1); _pgfont(run_pg)
    instr = run_pg._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = ' PAGE '
    run_pg2 = rule_para.add_run(); run_pg2._element.append(instr); _pgfont(run_pg2)
    fld2 = run_pg2._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_pg3 = rule_para.add_run(); run_pg3._element.append(fld2)

    rule_para.add_run('\t')
    run_series = rule_para.add_run('MoneyRules Series')
    run_series.font.size = Pt(7); run_series.font.italic = True; run_series.font.color.rgb = LIGHT_GREY
    _set_run_font(run_series)

    rule_para.add_run('\t')
    run_brand = rule_para.add_run('www.incomeonline.info')
    run_brand.font.size = Pt(8); run_brand.font.bold = True; run_brand.font.color.rgb = PINK
    _set_run_font(run_brand)


# ---------------------------------------------------------------
# Document factory
# ---------------------------------------------------------------

def _furnish_content_section(section, clean_first_page):
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    section.different_first_page_header_footer = clean_first_page
    _add_page_borders(section)
    _add_running_header(section)
    _add_footer(section)


def _add_full_page_cover(doc, section, image_path):
    """Turn `section` into a borderless, zero-margin page filled by the cover image."""
    section.top_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.right_margin = 0
    section.header_distance = 0
    section.footer_distance = 0
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(image_path, width=section.page_width,
                    height=int(section.page_height * 0.997))


def create_moneyrules_document(title='', subtitle='', cover_image=None):
    doc = Document()

    # Normal style: Montserrat 11pt, bold for strong on-screen/print legibility
    normal = doc.styles['Normal']
    _set_style_font(normal, FONT)
    normal.font.size = Pt(11)
    normal.font.bold = True
    normal.font.color.rgb = BODY_TEXT
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.32

    # Heading styles (level 2/3 — level 1 rendered as a branded panel)
    for level, (size, color, before) in {
        1: (Pt(16), DEEP_PURPLE, Pt(18)),
        2: (Pt(13), PINK, Pt(14)),
        3: (Pt(11.5), ORANGE, Pt(10)),
    }.items():
        hs = doc.styles[f'Heading {level}']
        _set_style_font(hs, FONT)
        hs.font.size = size
        hs.font.color.rgb = color
        hs.font.bold = True
        hs.paragraph_format.space_before = before
        hs.paragraph_format.space_after = Pt(6)
        hs.paragraph_format.keep_with_next = True

    # List styles → Montserrat, bold
    for sname in ('List Bullet', 'List Number'):
        try:
            _set_style_font(doc.styles[sname], FONT)
            doc.styles[sname].font.bold = True
        except Exception:
            pass

    if cover_image:
        # Section 0 = full-bleed cover; section 1 = furnished content.
        _add_full_page_cover(doc, doc.sections[0], cover_image)
        doc.add_section(WD_SECTION.NEW_PAGE)
        _furnish_content_section(doc.sections[-1], clean_first_page=False)
    else:
        for section in doc.sections:
            _furnish_content_section(section, clean_first_page=True)

    return doc


# ---------------------------------------------------------------
# Title / closing pages
# ---------------------------------------------------------------

def _rule(doc, position='bottom'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    pBdr.append(pBdr.makeelement(qn(f'w:{position}'), {
        qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): BORDER_OUTER,
    }))
    pPr.append(pBdr)
    return p


def add_title_page(doc, title, subtitle='', tagline=''):
    for _ in range(3):
        doc.add_paragraph()

    _rule(doc, 'bottom').paragraph_format.space_after = Pt(22)

    series = doc.add_paragraph()
    series.alignment = WD_ALIGN_PARAGRAPH.CENTER
    series.paragraph_format.space_after = Pt(4)
    r = series.add_run('M O N E Y R U L E S   S E R I E S')
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = PINK; _set_run_font(r)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(8)
    r = tp.add_run(title)
    r.font.size = Pt(34); r.bold = True; r.font.color.rgb = DEEP_PURPLE; _set_run_font(r)

    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(6)
        r = sp.add_run(subtitle)
        r.font.size = Pt(15); r.italic = True; r.font.color.rgb = ROSE; _set_run_font(r)

    if tagline:
        doc.add_paragraph()
        tg = doc.add_paragraph()
        tg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tg.paragraph_format.space_after = Pt(4)
        r = tg.add_run(tagline)
        r.font.size = Pt(10.5); r.italic = True; r.font.color.rgb = GREY; _set_run_font(r)

    _rule(doc, 'top')

    for _ in range(4):
        doc.add_paragraph()

    brand = doc.add_paragraph(); brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_after = Pt(2)
    r = brand.add_run('Brought to you by')
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GREY; _set_run_font(r)

    name = doc.add_paragraph(); name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    r = name.add_run('Income Online')
    r.font.size = Pt(15); r.bold = True; r.font.color.rgb = DEEP_PURPLE; _set_run_font(r)

    url = doc.add_paragraph(); url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = url.add_run('www.incomeonline.info')
    r.font.size = Pt(9); r.font.color.rgb = PINK; _set_run_font(r)

    doc.add_page_break()


def add_closing_page(doc):
    doc.add_paragraph()
    _rule(doc, 'bottom')
    doc.add_paragraph()

    final = doc.add_paragraph(); final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = final.add_run('Thank you for reading.')
    r.font.size = Pt(14); r.italic = True; r.font.color.rgb = DEEP_PURPLE; _set_run_font(r)

    doc.add_paragraph()
    bf = doc.add_paragraph(); bf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bf.paragraph_format.space_after = Pt(2)
    r = bf.add_run('Income Online')
    r.font.size = Pt(12); r.bold = True; r.font.color.rgb = DEEP_PURPLE; _set_run_font(r)

    tag = doc.add_paragraph(); tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_after = Pt(2)
    r = tag.add_run('Your Guide to Earning More')
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = PINK; _set_run_font(r)

    url = doc.add_paragraph(); url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url.paragraph_format.space_after = Pt(18)
    r = url.add_run('www.incomeonline.info')
    r.font.size = Pt(10); r.bold = True; r.font.color.rgb = PINK; _set_run_font(r)

    _rule(doc, 'top')
    doc.add_paragraph()

    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(
        '\u00A9 2026 Income Online. All Rights Reserved.\n'
        'This document is for educational purposes only and does not constitute financial advice.'
    )
    r.font.size = Pt(7.5); r.italic = True; r.font.color.rgb = LIGHT_GREY; _set_run_font(r)


# ---------------------------------------------------------------
# Headings — level 1 is a full-width branded purple panel
# ---------------------------------------------------------------

def add_styled_heading(doc, text, level=1):
    if level == 1:
        return _add_chapter_panel(doc, text)
    heading = doc.add_heading(text, level=level)
    for r in heading.runs:
        _set_run_font(r)
    return heading


def _add_chapter_panel(doc, text):
    """Branded full-width purple chapter heading panel (white text)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _full_width(table)
    _cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PURPLE_HEX)
    _cell_margins(cell, top=170, bottom=170, start=220, end=220)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(16); r.bold = True; r.font.color.rgb = WHITE; _set_run_font(r)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.keep_with_next = True
    return table


# ---------------------------------------------------------------
# Body + callouts
# ---------------------------------------------------------------

def add_body_text(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.36
    r = p.add_run(text)
    r.font.size = size or Pt(11)
    r.font.color.rgb = BODY_TEXT
    r.bold = True
    r.italic = italic
    _set_run_font(r)
    return p


def _callout(doc, title, body, accent_hex, tint_hex, title_color, bulleted=False):
    """
    Generic coloured callout: single-row table (cantSplit) so it never splits or
    orphans at the bottom of a page — the whole block moves to the next page.
    `body` may be a string or a list; list items become bullets when bulleted=True,
    otherwise separate prose paragraphs.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _full_width(table)
    _cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, tint_hex)
    _accent_left_border(cell, accent_hex, tint_hex)
    _cell_margins(cell)

    # Title line
    tp = cell.paragraphs[0]
    tp.paragraph_format.space_after = Pt(4)
    tp.paragraph_format.keep_with_next = True
    if title:
        tr = tp.add_run(title)
        tr.font.size = Pt(11); tr.bold = True; tr.font.color.rgb = title_color; _set_run_font(tr)

    items = body if isinstance(body, (list, tuple)) else [body]
    for i, item in enumerate(items):
        para = tp if (not title and i == 0) else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        prefix = '•  ' if bulleted else ''
        r = para.add_run(f'{prefix}{item}')
        r.font.size = Pt(10.5); r.bold = True; r.font.color.rgb = BODY_TEXT; _set_run_font(r)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_highlight_box(doc, text):
    """Branded key-takeaway / tip panel (purple). Kept together, never orphaned."""
    return _callout(doc, '', text, PURPLE_HEX, TINT_PURPLE, DEEP_PURPLE)


def add_expert_tip(doc, text, title='IncomeOnline Expert Tip'):
    return _callout(doc, title, text, PURPLE_HEX, TINT_PURPLE, DEEP_PURPLE)


def add_action_checklist(doc, items, title='Action Checklist'):
    return _callout(doc, title, list(items), PINK_HEX, TINT_PINK, PINK, bulleted=True)


def add_common_mistake(doc, text, title='Common Mistake'):
    return _callout(doc, title, text, ORANGE_HEX, TINT_ORANGE, ORANGE)


def add_example_box(doc, text, title='Example'):
    return _callout(doc, title, text, '7C3AED', TINT_PURPLE, MEDIUM_PURPLE)


def add_case_study(doc, text, title='IncomeOnline Case Study'):
    return _callout(doc, title, text, PINK_HEX, TINT_PINK, PINK)


# ---------------------------------------------------------------
# Tables
# ---------------------------------------------------------------

def add_branded_table(doc, headers, data, header_color='6D28D9'):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        _cant_split(row)

    # header row repeats across pages
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn('w:tblHeader'), {}))

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        set_cell_shading(cell, header_color)
        _set_cell_borders(cell, color=header_color, size='4')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
        r = p.add_run(header)
        r.font.size = Pt(9); r.bold = True; r.font.color.rgb = WHITE; _set_run_font(r)

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ''
            set_cell_shading(cell, TINT_PURPLE if row_idx % 2 == 0 else 'FFFFFF')
            _set_cell_borders(cell, color='E4D9FB', size='2')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(str(cell_text))
            r.font.size = Pt(9); r.bold = True; r.font.color.rgb = BODY_TEXT; _set_run_font(r)

    doc.add_paragraph()
    return table


def save_to_buffer(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
