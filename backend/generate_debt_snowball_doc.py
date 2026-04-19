"""Generate The Debt Snowball Method guide using the MoneyRules template."""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, PURPLE, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_debt_snowball_document():
    doc = create_moneyrules_document(
        title='The Debt Snowball Method',
        subtitle='How to Clear Your Debts in Record Time — Starting This Month'
    )
    add_title_page(doc,
        title='The Debt Snowball',
        subtitle='The Proven Method to Clear Your Debts Forever',
        tagline='Dave Ramsey\'s famous system, broken down\nstep by step, with worked UK examples.')

    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()
    for num, t in [
        ('1.', 'Introduction — Why Most People Stay in Debt'),
        ('2.', 'What Is the Debt Snowball?'),
        ('3.', 'Snowball vs Avalanche: Which Wins?'),
        ('4.', 'Step-by-Step: Build Your Snowball'),
        ('5.', 'A Worked Example with Real Numbers'),
        ('6.', 'The Psychology of Small Wins'),
        ('7.', 'When NOT to Use the Snowball'),
        ('8.', 'Dealing with Set-backs'),
        ('9.', 'Life After Debt — Keeping It Off'),
        ('10.', 'Key Takeaways and Your Action Plan'),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f'{num}  '); r.font.size = Pt(13); r.font.color.rgb = PURPLE; r.bold = True
        r2 = p.add_run(t); r2.font.size = Pt(13); r2.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run('DISCLAIMER: This guide is for educational purposes only. For serious debt problems, seek free advice from StepChange or Citizens Advice.')
    run.font.size = Pt(9); run.font.color.rgb = GREY; run.italic = True
    doc.add_page_break()

    add_styled_heading(doc, '1. Introduction — Why Most People Stay in Debt', level=1)
    add_body_text(doc, 'The average UK household in 2026 carries around £65,000 of debt, once mortgages are included. Exclude mortgages and it\'s still about £9,000 of credit cards, overdrafts and personal loans — each one slowly draining income as interest.')
    add_body_text(doc, 'Most people who try to clear their debts fail, not because they lack discipline, but because their strategy lacks momentum. They spread their repayments across every balance at once, never pay off a single debt in full, and give up.')
    add_highlight_box(doc, 'Debt clearance is a psychology problem first\nand a maths problem second.')
    doc.add_body_text if False else add_body_text(doc, 'The Debt Snowball fixes the psychology. The method is not the fastest arithmetically — but it is the fastest realistically, because humans actually stick with it.')
    doc.add_page_break()

    add_styled_heading(doc, '2. What Is the Debt Snowball?', level=1)
    add_body_text(doc, 'The Debt Snowball was popularised by US personal-finance author Dave Ramsey. The core rule is almost absurdly simple:')
    add_highlight_box(doc, '1. List your debts smallest to largest (ignore interest rate).\n2. Pay MINIMUMS on everything except the smallest.\n3. Throw every spare pound at the smallest debt.\n4. Once cleared, roll its payment into the next smallest.\n5. Repeat until debt-free.')
    add_body_text(doc, 'That is the whole method. No spreadsheets. No credit card transfers. No "consolidation loans" that take 7 years. Just a list, ordered by size, crushed from the top down.')
    doc.add_page_break()

    add_styled_heading(doc, '3. Snowball vs Avalanche: Which Wins?', level=1)
    add_body_text(doc, 'The rival method is the Debt Avalanche: instead of ordering by size, order by INTEREST RATE (highest first). Mathematically the Avalanche wins — slightly. So why does Snowball work better in practice?')
    add_branded_table(doc,
        headers=['Method', 'Order By', 'Mathematically', 'Behaviourally'],
        data=[
            ('Debt Snowball',  'Smallest balance first',  'Loses by ~2–5%',  'Wins big — fast early victories'),
            ('Debt Avalanche', 'Highest interest first',  'Wins by ~2–5%',   'Loses — no early wins to reinforce'),
        ])
    add_body_text(doc, 'Research from Northwestern University (Gal & McShane, 2012) showed Snowball users were significantly more likely to fully clear their debts. The psychology of banking an early win, then another, keeps you going for the long haul.')
    add_highlight_box(doc, 'A method you will complete beats a method that is 3% better on paper.')
    doc.add_page_break()

    add_styled_heading(doc, '4. Step-by-Step: Build Your Snowball', level=1)
    for head, body in [
        ('Step 1: List Every Debt',
         'Credit cards, overdraft, personal loans, store cards, car finance — every single one. For each, note: (a) balance, (b) minimum monthly payment, (c) interest rate (for reference only).'),
        ('Step 2: Order by Balance, Smallest First',
         'Ignore interest rate. Literally sort your list smallest balance on top, largest on bottom.'),
        ('Step 3: Pay Minimum on All BUT the Top',
         'This keeps you in good standing with every creditor. Missed payments damage credit scores and trigger late fees.'),
        ('Step 4: Every Spare Pound Goes to the Top Debt',
         'Side-hustle money, sold-on-eBay money, tax rebate, birthday cash — anything extra flows to the smallest balance until it is £0.'),
        ('Step 5: Roll & Repeat',
         'When the smallest debt clears, take the payment you were making on it and ADD it to the minimum of the next-smallest. The snowball grows every time.'),
    ]:
        add_styled_heading(doc, head, level=2)
        add_body_text(doc, body)
    doc.add_page_break()

    add_styled_heading(doc, '5. A Worked Example with Real Numbers', level=1)
    add_body_text(doc, 'Sarah has four debts and £250/month available for extra repayment above minimums:')
    add_branded_table(doc,
        headers=['Debt', 'Balance', 'Min/mo', 'APR'],
        data=[
            ('Store card',          '£380',    '£25',  '29.9%'),
            ('Credit card A',       '£1,200',  '£40',  '19.9%'),
            ('Overdraft',           '£1,500',  '£30',  '35.0%'),
            ('Personal loan',       '£5,400',  '£180', '12.5%'),
        ])
    add_body_text(doc, 'Month 1: Pay minimums (£275 total) + £250 extra onto the £380 store card = £275 paid on store card. Balance drops to £105.')
    add_body_text(doc, 'Month 2: Store card cleared (£105 + rolled £25 min = first win!). Spare cash now £275 extra. Target credit card A.')
    add_body_text(doc, 'Month 5: Credit card A cleared. Snowball is now £315 of monthly momentum. Target overdraft.')
    add_body_text(doc, 'Month 10: Overdraft cleared. £345/month snowball now hammers the personal loan.')
    add_body_text(doc, 'Month ~25: Personal loan done. Total debt-free in just over 2 years, despite starting with £8,480 of debt.')
    add_highlight_box(doc, '£250/month + the snowball effect\n= £8,480 cleared in 25 months.')
    doc.add_page_break()

    add_styled_heading(doc, '6. The Psychology of Small Wins', level=1)
    add_body_text(doc, 'Behavioural economists have long known that humans are terrible at pursuing distant rewards. A goal 24 months away feels abstract; a goal 2 weeks away feels real. The Snowball exploits this by engineering your first "debt cleared!" moment within 30–60 days for most people.')
    add_body_text(doc, 'That first win triggers a small dopamine release. Then another. Then another. Each win reinforces the habit. By the time you reach the big, scary debts, the routine is automatic.')
    add_highlight_box(doc, 'Motivation runs out.\nHabits run on autopilot.\nThe Snowball builds the habit.')
    doc.add_page_break()

    add_styled_heading(doc, '7. When NOT to Use the Snowball', level=1)
    add_body_text(doc, 'The Snowball is not universal. Skip it or modify it in these cases:')
    for t in [
        'You have debt with APR above 30% on a large balance — the extra interest cost of Snowball vs Avalanche becomes material. Switch to Avalanche.',
        'You have payday loans or illegal-lender debt — seek IMMEDIATE free help from StepChange or Citizens Advice before anything else.',
        'Your total debt exceeds 12 months of take-home pay — you may need a DMP (Debt Management Plan) or IVA rather than self-managed repayment.',
        'You have no emergency fund at all — build a £500–£1,000 starter fund first, otherwise the next car breakdown resets the whole plan.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    doc.add_page_break()

    add_styled_heading(doc, '8. Dealing with Set-backs', level=1)
    add_body_text(doc, 'Life happens. A boiler breaks, a car dies, a medical bill lands. Plan for this in advance rather than pretending it won\'t occur.')
    add_styled_heading(doc, 'The £1,000 Buffer', level=2)
    add_body_text(doc, 'Before starting the Snowball in earnest, park £500–£1,000 in an easy-access savings account. This buffer absorbs the ordinary shocks of life so one setback doesn\'t force you back to credit cards.')
    add_styled_heading(doc, 'Missed a Month? Don\'t Quit.', level=2)
    add_body_text(doc, 'A single missed extra payment is a blip, not a disaster. Make the minimums anyway, then restart the snowball next month exactly where you left off. The only way to truly fail is to abandon the plan entirely.')
    doc.add_page_break()

    add_styled_heading(doc, '9. Life After Debt — Keeping It Off', level=1)
    add_body_text(doc, 'Roughly 70% of people who clear their debts return to debt within 3 years. The Snowball gets you out; habits keep you out.')
    for t in [
        'Convert the snowball payment into a savings contribution — same money, opposite direction.',
        'Build a 3-month emergency fund (see our separate guide on this).',
        'Pay credit cards IN FULL every month — set a direct debit for the statement balance.',
        'If you cannot pay cash for it, do not buy it — with one exception (your home).',
        'Review budgets quarterly. Small lifestyle creep is how most people quietly re-enter debt.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_highlight_box(doc, 'The Snowball is a one-off event.\nStaying debt-free is a daily habit.')
    doc.add_page_break()

    add_styled_heading(doc, '10. Key Takeaways and Your Action Plan', level=1)
    for t in [
        'List every debt, ordered smallest balance first.',
        'Pay minimums on all, extras on the top.',
        'Bank the first win within 60 days — that\'s where the magic starts.',
        'Roll cleared payments into the next debt; the snowball grows.',
        'Build a £500–£1,000 emergency buffer BEFORE attacking.',
        'Avalanche wins on paper; Snowball wins in real life.',
        'After debt-free: redirect the snowball into savings/investments.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT

    add_styled_heading(doc, 'Your First Week', level=2)
    for t in [
        'Day 1: Write down every debt — name, balance, min, APR.',
        'Day 2: Sort by balance, smallest on top.',
        'Day 3: Open a £500–£1,000 starter emergency fund (savings account).',
        'Day 4: Calculate your "spare" £ for the snowball each month.',
        'Day 5: Set direct debits for all minimums.',
        'Day 6: Schedule the first extra payment for next payday.',
        'Day 7: Tell one trusted person. Accountability doubles your odds.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_body_text(doc, 'Visit www.incomeonline.info for 199+ legitimate ways to boost your snowball — every extra £50/month chops months off your payoff date.')
    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    buf = generate_debt_snowball_document()
    os.makedirs('/app/backend/static', exist_ok=True)
    with open('/app/backend/static/The_Debt_Snowball_Method.docx', 'wb') as f:
        f.write(buf.read())
    print('Debt Snowball guide generated!')
