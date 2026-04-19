"""Generate the Compound Interest Guide using the MoneyRules template."""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, DEEP_PURPLE, PURPLE, PINK, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_compound_interest_document():
    doc = create_moneyrules_document(
        title='The Compound Interest Handbook',
        subtitle='How Small Sums Become Large Fortunes — Mathematically Proven'
    )
    add_title_page(doc,
        title='Compound Interest',
        subtitle='The Eighth Wonder of the World',
        tagline='The formula, the psychology, and worked scenarios\nfor retirement, children\'s savings and house deposits.')

    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()
    for num, t in [
        ('1.', 'Introduction — Why Einstein Called It a Wonder'),
        ('2.', 'The Formula, Broken Down'),
        ('3.', 'Simple vs Compound: The Critical Difference'),
        ('4.', 'The Power of TIME (Not Amount)'),
        ('5.', 'The Effect of RATE: Why 1% Matters'),
        ('6.', 'The Effect of FREQUENCY (Annual/Monthly/Daily)'),
        ('7.', 'The Reverse Calculator: Saving for a Goal'),
        ('8.', 'Real Scenarios: Retirement, Kids, House Deposit'),
        ('9.', 'The Dark Side: When Compound Interest Works Against You'),
        ('10.', 'Key Takeaways and Action Plan'),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f'{num}  '); r.font.size = Pt(13); r.font.color.rgb = PURPLE; r.bold = True
        r2 = p.add_run(t); r2.font.size = Pt(13); r2.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run('DISCLAIMER: Educational content only. Not financial advice.')
    run.font.size = Pt(9); run.font.color.rgb = GREY; run.italic = True
    doc.add_page_break()

    add_styled_heading(doc, '1. Introduction — Why Einstein Called It a Wonder', level=1)
    add_body_text(doc, '"Compound interest is the eighth wonder of the world. He who understands it, earns it. He who doesn\'t, pays it." The quote is usually attributed to Einstein, though historians are unsure. What is certain is that compound interest is the most important single concept in personal finance — yet most people profoundly underestimate its power.')
    add_body_text(doc, 'The reason it feels magical is that it is EXPONENTIAL, and humans are evolved to think linearly. A savings account paying 7% per year doesn\'t grow 7% every year from the original balance — it grows 7% of the ever-larger total, including past interest. Each year\'s interest itself starts earning interest.')
    add_highlight_box(doc, 'Simple interest grows in a line.\nCompound interest grows in a curve.\nOver decades, the curve wins by miles.')
    doc.add_page_break()

    add_styled_heading(doc, '2. The Formula, Broken Down', level=1)
    add_body_text(doc, 'The compound interest formula is:')
    f1 = doc.add_paragraph(); f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f1.paragraph_format.space_before = Pt(12); f1.paragraph_format.space_after = Pt(12)
    run = f1.add_run('A = P × (1 + r/n)^(n·t)')
    run.font.size = Pt(18); run.font.color.rgb = DEEP_PURPLE; run.bold = True
    add_body_text(doc, 'Where:')
    for t in ['A = Final amount', 'P = Principal (starting amount)', 'r = Annual interest rate (as a decimal, e.g. 0.05 for 5%)', 'n = Compoundings per year (12 = monthly, 365 = daily)', 't = Time in years']:
        p = doc.add_paragraph(t, style='List Bullet')
        for rn in p.runs: rn.font.size = Pt(11); rn.font.color.rgb = BODY_TEXT
    add_styled_heading(doc, 'A Manual Calculation', level=2)
    add_body_text(doc, 'You deposit £10,000 at 6% annually, compounded monthly, for 20 years. A = 10,000 × (1 + 0.06/12)^(12 × 20) = 10,000 × (1.005)^240 ≈ £33,102.')
    add_body_text(doc, 'You put in £10,000. You end with £33,102. Over £23,000 of pure growth, from doing nothing but waiting.')
    doc.add_page_break()

    add_styled_heading(doc, '3. Simple vs Compound: The Critical Difference', level=1)
    add_body_text(doc, 'Consider £10,000 invested at 7% for 40 years:')
    add_branded_table(doc,
        headers=['Method', 'Final Value', 'Growth'],
        data=[
            ('Simple interest', '£38,000',   '3.8× starting amount'),
            ('Compound interest (annual)', '£149,744',  '14.9× starting amount'),
        ])
    add_body_text(doc, 'Same money. Same rate. Same time period. Four times the final wealth — purely because compound re-invests the interest to earn more interest.')
    add_highlight_box(doc, 'The difference between simple and compound\nis the difference between middle-class and wealthy.')
    doc.add_page_break()

    add_styled_heading(doc, '4. The Power of TIME (Not Amount)', level=1)
    add_body_text(doc, 'This may be the single most important financial chart of your life. Two investors, both retire at 65. They earn the same 7% annual return.')
    add_branded_table(doc,
        headers=['Investor', 'Saves from age…', 'Total invested', 'Value at 65'],
        data=[
            ('Anna',   '25 to 35 (10 yrs), then stops',  '£24,000 (£200/mo × 10 yrs)', '£266,000'),
            ('Ben',    '35 to 65 (30 yrs)',              '£72,000 (£200/mo × 30 yrs)', '£245,000'),
        ])
    add_body_text(doc, 'Anna invested only £24,000 and stopped. Ben invested THREE TIMES as much, for THREE TIMES as long. Anna still ends up with more.')
    add_body_text(doc, 'Why? Anna\'s money had 10 extra years of compounding. Those 10 years are worth more than the next 30 combined.')
    add_highlight_box(doc, 'Starting early beats investing more.\nEvery single time.')
    doc.add_page_break()

    add_styled_heading(doc, '5. The Effect of RATE: Why 1% Matters', level=1)
    add_body_text(doc, '£10,000 invested for 30 years at different rates:')
    add_branded_table(doc,
        headers=['Rate', 'Final Value', 'Extra vs 4%'],
        data=[
            ('3%',  '£24,273',   '—'),
            ('4%',  '£32,434',   'Baseline'),
            ('5%',  '£43,219',   '+£10,785'),
            ('6%',  '£57,435',   '+£25,001'),
            ('7%',  '£76,123',   '+£43,689'),
            ('8%',  '£100,627',  '+£68,193'),
        ])
    add_body_text(doc, 'A 1% higher rate over 30 years adds £10,000+ on just £10k starting capital. This is why paying a 1.5% fund management fee instead of 0.2% is catastrophic over a career.')
    add_highlight_box(doc, 'Low-cost index funds keep your 1% in your pocket.\nOver 30 years, that 1% becomes hundreds of thousands.')
    doc.add_page_break()

    add_styled_heading(doc, '6. The Effect of FREQUENCY (Annual/Monthly/Daily)', level=1)
    add_body_text(doc, '£10,000 at 6% for 20 years, compounded at different intervals:')
    add_branded_table(doc,
        headers=['Compounding', 'Formula Effect', 'Final Value'],
        data=[
            ('Annually (n=1)',    '(1.06)^20',           '£32,071'),
            ('Quarterly (n=4)',   '(1.015)^80',          '£32,907'),
            ('Monthly (n=12)',    '(1.005)^240',         '£33,102'),
            ('Daily (n=365)',     '(1.000164)^7,300',    '£33,198'),
            ('Continuous',         'e^(0.06 × 20)',      '£33,201'),
        ],
        header_color='EA580C')
    add_body_text(doc, 'More-frequent compounding helps, but with diminishing returns. Moving from annual to monthly is worth about £1,000 here. Moving from monthly to daily is worth only £100 more. Don\'t obsess over this — rate and time matter far more.')
    doc.add_page_break()

    add_styled_heading(doc, '7. The Reverse Calculator: Saving for a Goal', level=1)
    add_body_text(doc, 'Often the useful question is not "how much will I have?" but "how much should I save?"')
    add_styled_heading(doc, 'Formula for Monthly Saving Needed', level=2)
    f2 = doc.add_paragraph(); f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f2.add_run('PMT = FV × r / ( (1+r)^n − 1 )')
    run.font.size = Pt(14); run.font.color.rgb = PINK; run.bold = True
    add_body_text(doc, 'Where FV = target value, r = monthly rate, n = number of months.')
    add_styled_heading(doc, 'Worked Example', level=2)
    add_body_text(doc, 'You want £100,000 in 20 years. Assumed return: 7% annual (0.583% monthly).')
    add_body_text(doc, 'PMT = 100,000 × 0.00583 / ((1.00583)^240 − 1) ≈ £192/month')
    add_highlight_box(doc, '£192/month for 20 years → £100,000.\nYou\'ve contributed £46,000 of it. Compounding added £54,000.')
    doc.add_page_break()

    add_styled_heading(doc, '8. Real Scenarios: Retirement, Kids, House Deposit', level=1)
    add_styled_heading(doc, 'Retirement at 65 starting at 30', level=2)
    add_body_text(doc, '£300/month into a pension at 7% from age 30 to 65 = £543,000. The same £300/month from age 45 to 65 = only £156,000. 15 extra years = 3.5× more wealth.')
    add_styled_heading(doc, 'University fund for a newborn', level=2)
    add_body_text(doc, '£100/month in a Junior ISA at 6% for 18 years = £38,929. £21,600 of that is your contribution — £17,329 is pure compounding growth.')
    add_styled_heading(doc, 'House deposit in 7 years', level=2)
    add_body_text(doc, 'Target £30,000 in 7 years at 5% = £300/month. At 3% (cash savings only) you\'d need £325/month. A Lifetime ISA with its 25% government bonus turns £300/month into your full £30,000 in ~5 years.')
    add_highlight_box(doc, 'Pair compounding with government-boosted accounts\n(Pension relief, LISA bonus)\n— free money on top of already-exponential growth.')
    doc.add_page_break()

    add_styled_heading(doc, '9. The Dark Side: When Compound Interest Works Against You', level=1)
    add_body_text(doc, 'The same mathematical force that grows your wealth grows your debt — only faster, because consumer debt rates are typically far higher than investment returns.')
    add_branded_table(doc,
        headers=['Debt Type', 'Typical APR', '£1,000 after 10 yrs unpaid'],
        data=[
            ('Mortgage',             '4–6%',    '£1,480 – £1,790'),
            ('Personal loan',        '8–15%',   '£2,160 – £4,046'),
            ('Credit card',          '20–30%',  '£6,192 – £13,786'),
            ('Payday loan',          '400%+',   'Catastrophic'),
        ])
    add_body_text(doc, 'Minimum-payment credit card debt is financial quicksand. The bank is compounding YOUR money into THEIR pocket, often at rates 3–5× faster than your investments can grow. Kill high-interest debt before investing.')
    add_highlight_box(doc, 'The question isn\'t whether compound interest will shape your life.\nIt\'s whether it will shape it for you or against you.')
    doc.add_page_break()

    add_styled_heading(doc, '10. Key Takeaways and Action Plan', level=1)
    for t in [
        'Compound = interest earning interest. It\'s exponential, not linear.',
        'TIME matters more than AMOUNT — start as early as possible.',
        '1% more return can mean hundreds of thousands of pounds over a career.',
        'Monthly/daily compounding is nice but not critical. Rate + time are what matter.',
        'Use the reverse formula to work out how much to save for a goal.',
        'Clear high-interest debt FIRST — it compounds against you faster than investments compound for you.',
        'Pair compounding with tax wrappers (ISA, SIPP, LISA) for double benefit.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT

    add_styled_heading(doc, 'Your First Week', level=2)
    for t in [
        'Day 1: Calculate what £100/month at 7% for YOUR remaining years to 65 would produce.',
        'Day 2: Open a Stocks & Shares ISA if you don\'t have one (~10 minutes).',
        'Day 3: Set up a monthly direct debit — even £50 is a start.',
        'Day 4: Review any credit card debt above 15% APR — attack this first.',
        'Day 5: Pick an ACCUMULATION share class (re-invests dividends automatically).',
        'Day 6: Diarise a quarterly review to increase contributions as income rises.',
        'Day 7: Leave it alone. Seriously. The magic happens in the decades, not days.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_body_text(doc, 'Visit www.incomeonline.info for 199+ ways to grow the "invest more each month" side of the equation.')
    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    buf = generate_compound_interest_document()
    os.makedirs('/app/backend/static', exist_ok=True)
    with open('/app/backend/static/Compound_Interest_Handbook.docx', 'wb') as f:
        f.write(buf.read())
    print('Compound Interest guide generated!')
