"""Three compact MoneyRules-style guides sharing a single module for efficiency."""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, PURPLE, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _toc(doc, items):
    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()
    for num, t in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f'{num}  '); r.font.size = Pt(13); r.font.color.rgb = PURPLE; r.bold = True
        r2 = p.add_run(t); r2.font.size = Pt(13); r2.font.color.rgb = BODY_TEXT
    doc.add_paragraph()
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = d.add_run('DISCLAIMER: Educational content only. Not financial advice. Reflects 2026 UK rules.')
    run.font.size = Pt(9); run.font.color.rgb = GREY; run.italic = True
    doc.add_page_break()


def _bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT


def _numbers(doc, items):
    for t in items:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT


# =====================================================================
# GUIDE 8 — UK Credit Score Masterclass
# =====================================================================
def generate_credit_score_document():
    doc = create_moneyrules_document(
        title='UK Credit Score Masterclass',
        subtitle='Understand, Build and Repair Your Credit — The Complete Guide'
    )
    add_title_page(doc,
        title='Your Credit Score',
        subtitle='Master It, Maintain It, Leverage It',
        tagline='Why it matters, how it\'s calculated,\nand 7 proven ways to lift it in 90 days.')
    _toc(doc, [
        ('1.', 'Introduction — Why Your Score Controls Your Life'),
        ('2.', 'The Three UK Credit Bureaus'),
        ('3.', 'What Actually Moves Your Score'),
        ('4.', '7 Proven Ways to Lift Your Score in 90 Days'),
        ('5.', 'Common Credit Mistakes'),
        ('6.', 'Repairing a Damaged Score'),
        ('7.', 'How Lenders Really Read Your Report'),
        ('8.', 'Your Action Plan'),
    ])

    add_styled_heading(doc, '1. Introduction — Why Your Score Controls Your Life', level=1)
    add_body_text(doc, 'Your credit score quietly determines a huge portion of your financial life. Mortgage rates, credit card limits, car finance, phone contracts, rental applications, even some utility deposits — all are decided by a three-digit number most people never actively manage.')
    add_body_text(doc, 'The gap between a "poor" score (300-500) and an "excellent" score (800+) can easily be £50,000+ across a mortgage lifetime. Good news: the score is built from your behaviour, which means you can improve it with discipline and patience.')
    add_highlight_box(doc, 'Your credit score is the single highest-ROI\n3-digit number you\'ll ever track.')
    doc.add_page_break()

    add_styled_heading(doc, '2. The Three UK Credit Bureaus', level=1)
    add_body_text(doc, 'The UK has three major credit reference agencies. Each uses a slightly different scoring scale but pulls from similar data:')
    add_branded_table(doc,
        headers=['Bureau', 'Score Range', 'Free Service'],
        data=[
            ('Experian',   '0 – 999',   'Get a free report via MSE Credit Club'),
            ('Equifax',    '0 – 1,000', 'Free via ClearScore'),
            ('TransUnion', '0 – 710',   'Free via Credit Karma'),
        ])
    add_body_text(doc, 'Lenders typically check one or more of these. Your scores will differ slightly across bureaus because each holds slightly different data. Never assume the one you checked is what the lender saw.')
    doc.add_page_break()

    add_styled_heading(doc, '3. What Actually Moves Your Score', level=1)
    add_branded_table(doc,
        headers=['Factor', 'Weight', 'What It Means'],
        data=[
            ('Payment history',      '35%', 'On-time vs late payments in last 6 years'),
            ('Credit utilisation',   '30%', 'How much of your available credit you use'),
            ('Length of history',    '15%', 'How long your oldest account has been open'),
            ('New applications',     '10%', 'Recent hard searches damage the score short-term'),
            ('Credit mix',           '10%', 'Mix of credit cards, loans, mortgages'),
        ])
    add_highlight_box(doc, 'Payment history and utilisation alone = 65%.\nGet these two right and everything else is noise.')
    doc.add_page_break()

    add_styled_heading(doc, '4. 7 Proven Ways to Lift Your Score in 90 Days', level=1)
    _numbers(doc, [
        'Pay every bill by direct debit — missed payments are the #1 score killer.',
        'Keep credit utilisation below 30%. Under 10% is ideal.',
        'Register on the electoral roll — cheapest, fastest boost available.',
        'Check your file and dispute any errors (about 20% of reports have errors).',
        'Don\'t close old credit cards — closing shortens history and raises utilisation.',
        'Avoid new applications 3 months before a big decision (mortgage/car).',
        'Link your current account to credit reporting services like Experian Boost.',
    ])
    doc.add_page_break()

    add_styled_heading(doc, '5. Common Credit Mistakes', level=1)
    _bullets(doc, [
        'Carrying credit cards at 80%+ utilisation — crushes your score even when paid.',
        'Applying for multiple cards in quick succession — each triggers a hard search.',
        'Closing paid-off credit cards — kills your credit history length and raises utilisation.',
        'Paying only the minimum on credit cards — costs you massive interest + keeps balance high.',
        'Ignoring old addresses on your file — can confuse identity verification.',
        'Not having ANY credit — counter-intuitive, but zero history = lenders can\'t assess you.',
    ])
    doc.add_page_break()

    add_styled_heading(doc, '6. Repairing a Damaged Score', level=1)
    add_body_text(doc, 'A damaged score is not a life sentence. Most negative entries age off your file after 6 years. Your immediate mission is to stop the bleeding and start rebuilding:')
    _numbers(doc, [
        'Immediately bring all accounts current. One missed payment damages more than a dozen late ones.',
        'Consider a credit-builder card (Aqua, Vanquis, Capital One). High APR, low limit — use 5-15% of limit + pay in full monthly.',
        'Avoid applying for ANY new credit for 6 months while the score rebuilds.',
        'Dispute factual errors aggressively — lenders must respond within 28 days.',
        'Time + consistency does the rest. Expect 50-150 point recovery over 12-18 months.',
    ])
    add_highlight_box(doc, 'You can\'t buy credit repair. Anyone promising to\n"fix" your score for a fee is running a scam.')
    doc.add_page_break()

    add_styled_heading(doc, '7. How Lenders Really Read Your Report', level=1)
    add_body_text(doc, 'Scores are a summary. Lenders also look at the raw report, applying their own affordability model. They care about:')
    _bullets(doc, [
        'Stability — same address, same employer for 2+ years is a huge positive signal.',
        'Debt-to-income ratio — total monthly debts under 40% of net income.',
        'Spending patterns — consistent outgoings vs volatile swings.',
        'Gambling or short-term lending usage — red flags in the bank statement review.',
        'Employment type — PAYE employment is weighted more favourably than self-employed.',
    ])
    doc.add_page_break()

    add_styled_heading(doc, '8. Your Action Plan', level=1)
    _numbers(doc, [
        'Week 1: Sign up to all 3 free bureau services and download your reports.',
        'Week 2: Register (or confirm) on the electoral roll in ALL current addresses.',
        'Week 3: Set direct debits for every single bill and minimum credit card payment.',
        'Week 4: Review utilisation on every card — pay down to under 30% of limit.',
        'Month 2: Dispute any errors; wait for corrections.',
        'Month 3: Re-check scores. Most people see a 20-60 point lift by now.',
    ])
    add_body_text(doc, 'Visit www.incomeonline.info to discover 199+ platforms that can supplement your income — the surest way to improve credit utilisation is to have more income to work with.')
    add_closing_page(doc)
    return save_to_buffer(doc)


# =====================================================================
# GUIDE 9 — ISA vs SIPP
# =====================================================================
def generate_isa_vs_sipp_document():
    doc = create_moneyrules_document(
        title='ISA vs SIPP — The Complete UK Tax-Efficient Investing Guide',
        subtitle='Which Account, When, and Why — Decoded Without the Jargon'
    )
    add_title_page(doc,
        title='ISA vs SIPP',
        subtitle='The UK Investor\'s Complete Comparison',
        tagline='Flexibility, tax relief, lock-ins and limits\n— explained so you pick correctly first time.')
    _toc(doc, [
        ('1.', 'Introduction — Why the UK\'s Two Tax Wrappers Matter'),
        ('2.', 'ISA: The Flexible Tax-Free Account'),
        ('3.', 'SIPP: The Pension Tax Turbocharger'),
        ('4.', 'Head-to-Head Comparison'),
        ('5.', 'Which to Use When — 5 Life Scenarios'),
        ('6.', 'The LISA and Other Special Cases'),
        ('7.', 'Tax Reliefs Worked Through with Numbers'),
        ('8.', 'Your Action Plan'),
    ])

    add_styled_heading(doc, '1. Introduction — Why the UK\'s Two Tax Wrappers Matter', level=1)
    add_body_text(doc, 'The UK government offers two powerful tax shelters for long-term investing: Individual Savings Accounts (ISAs) and Self-Invested Personal Pensions (SIPPs). If you\'re not using at least one, you\'re handing HMRC money you don\'t need to.')
    add_highlight_box(doc, 'Every pound outside these wrappers pays tax twice:\nincome tax on earnings + capital gains tax on growth.')
    doc.add_page_break()

    add_styled_heading(doc, '2. ISA: The Flexible Tax-Free Account', level=1)
    add_body_text(doc, 'An ISA is a container. Whatever you put inside grows tax-free — no income tax on dividends, no capital gains tax on growth, no tax on withdrawal. Ever.')
    add_body_text(doc, '2026 allowance: £20,000 per adult per tax year (6 April – 5 April). Couples can shelter £40,000 combined annually.')
    add_styled_heading(doc, 'ISA Types', level=2)
    add_branded_table(doc,
        headers=['Type', 'Purpose', 'Notes'],
        data=[
            ('Cash ISA',              'Savings (bank-like)',  '4-5% rates in 2026'),
            ('Stocks & Shares ISA',   'Investing',            'Most flexible for long-term growth'),
            ('Lifetime ISA (LISA)',   'First home or retirement', '25% govt bonus, age 18-39 only'),
            ('Innovative Finance ISA','P2P lending',          'Higher risk, usually avoid'),
        ])
    doc.add_page_break()

    add_styled_heading(doc, '3. SIPP: The Pension Tax Turbocharger', level=1)
    add_body_text(doc, 'A Self-Invested Personal Pension is a DIY pension. Every contribution receives tax relief from HMRC — effectively a 20%, 40% or 45% top-up depending on your tax band.')
    add_body_text(doc, '2026 annual allowance: £60,000 or 100% of earnings (whichever is lower). Money is locked in until age 57 (rising to 58 by 2028).')
    add_styled_heading(doc, 'How the Tax Relief Works', level=2)
    add_branded_table(doc,
        headers=['Tax Band', 'Your Contribution', 'HMRC Tops Up', 'In Your SIPP'],
        data=[
            ('Basic-rate (20%)',   '£80',  '£20',  '£100'),
            ('Higher-rate (40%)',  '£60',  '£40',  '£100'),
            ('Additional (45%)',   '£55',  '£45',  '£100'),
        ])
    add_highlight_box(doc, 'A higher-rate taxpayer gets £40 of free money\nfor every £60 they put into a SIPP.')
    doc.add_page_break()

    add_styled_heading(doc, '4. Head-to-Head Comparison', level=1)
    add_branded_table(doc,
        headers=['Feature', 'ISA', 'SIPP'],
        data=[
            ('Annual allowance',        '£20,000',         '£60,000 (or 100% earnings)'),
            ('Tax relief on deposit',   'None',            '20/40/45% top-up'),
            ('Tax on growth',           'None',            'None'),
            ('Tax on withdrawal',       'None',            '25% tax-free, rest at your income tax rate'),
            ('Access age',              'Any time',        '57+'),
            ('Inheritance',             'Part of estate',  'Usually outside estate'),
        ])
    doc.add_page_break()

    add_styled_heading(doc, '5. Which to Use When — 5 Life Scenarios', level=1)
    for h, b in [
        ('Young professional (20s-30s)',
         'Max LISA first (£4,000/yr = £1,000 free bonus), then ISA. Start a small SIPP once higher-rate tax kicks in.'),
        ('Higher-rate taxpayer',
         'SIPP first, hard. The 40% tax relief is uniquely powerful — more than makes up for the age lock.'),
        ('Saving for a house',
         'LISA + cash ISA combo. LISA bonus = 25% free money on up to £4,000/yr towards a first home.'),
        ('Near retirement (50s-60s)',
         'Mix both. SIPP for the tax relief on what\'s left of your career; ISA for flexibility in early retirement.'),
        ('Self-employed',
         'SIPP is especially valuable — no employer pension means no workplace contributions, so DIY matters more.'),
    ]:
        add_styled_heading(doc, h, level=2)
        add_body_text(doc, b)
    doc.add_page_break()

    add_styled_heading(doc, '6. The LISA and Other Special Cases', level=1)
    add_body_text(doc, 'The Lifetime ISA (LISA) deserves its own section. It combines ISA tax-free growth with a 25% government bonus on contributions up to £4,000/year.')
    _bullets(doc, [
        'Must open between age 18-39.',
        'Maximum contribution £4,000/year.',
        'Government adds 25% bonus — £1,000/year free.',
        'Money can be used for first home purchase OR from age 60 onwards.',
        'Early withdrawal for any other reason = 25% penalty (loses your bonus + more).',
        'Counts against your £20,000 ISA allowance.',
    ])
    add_highlight_box(doc, 'If you\'re under 40 and buying your first home,\nthe LISA is almost always your best move FIRST.')
    doc.add_page_break()

    add_styled_heading(doc, '7. Tax Reliefs Worked Through with Numbers', level=1)
    add_body_text(doc, 'Consider Alice, a higher-rate taxpayer who wants to invest £500/month for 25 years at 7% return:')
    add_branded_table(doc,
        headers=['Wrapper', 'Her £ Contribution', 'Effective £ Invested', 'Final Value', 'Net After Tax'],
        data=[
            ('No wrapper',    '£150,000', '£150,000', '£394,000', '~£316,000 (after CGT)'),
            ('ISA',           '£150,000', '£150,000', '£394,000', '£394,000 (tax-free)'),
            ('SIPP',          '£150,000', '£250,000',  '£658,000', '~£542,000 (after pension tax)'),
        ])
    add_body_text(doc, 'SIPP wins by ~£148,000 vs ISA in this scenario purely from the 40% tax relief compounding over 25 years. But that money is locked until 57.')
    doc.add_page_break()

    add_styled_heading(doc, '8. Your Action Plan', level=1)
    _numbers(doc, [
        'Day 1: Check your tax band — basic, higher, or additional.',
        'Day 2: Open a Stocks & Shares ISA at a low-cost broker (Vanguard, InvestEngine, AJ Bell).',
        'Day 3: If under 40 and saving for a home — open a LISA.',
        'Day 4: If higher-rate taxpayer — open a SIPP alongside the ISA.',
        'Day 5: Set up direct debits to both, split based on your goals.',
        'Day 6: Choose global index funds (VWRL, VUSA, or FTSE Global All Cap).',
        'Day 7: Leave it alone. Decades of compounding > any clever trading.',
    ])
    add_body_text(doc, 'Visit www.incomeonline.info for 199+ ways to grow the income side of your equation — extra income = extra ISA/SIPP contributions = exponential long-term wealth.')
    add_closing_page(doc)
    return save_to_buffer(doc)


# =====================================================================
# GUIDE 10 — Side-Hustle Quick-Start Guide
# =====================================================================
def generate_side_hustle_document():
    doc = create_moneyrules_document(
        title='The Side-Hustle Quick-Start Guide',
        subtitle='Launch a Paying Side Business in 30 Days — Even From Your Sofa'
    )
    add_title_page(doc,
        title='Side-Hustle Quick-Start',
        subtitle='From Zero to First Payment in 30 Days',
        tagline='The 7 best side-hustle categories, how to pick yours,\nand a day-by-day launch plan.')
    _toc(doc, [
        ('1.', 'Introduction — Why Every Adult Needs a Side-Hustle in 2026'),
        ('2.', 'The 7 Side-Hustle Categories (with Real Earnings)'),
        ('3.', 'How to Pick YOUR Side-Hustle in 15 Minutes'),
        ('4.', 'The 30-Day Launch Plan'),
        ('5.', 'Finding Your First 3 Customers'),
        ('6.', 'Common Traps to Avoid'),
        ('7.', 'Scaling From £200 to £2,000 a Month'),
        ('8.', 'Your Action Plan'),
    ])

    add_styled_heading(doc, '1. Introduction — Why Every Adult Needs a Side-Hustle in 2026', level=1)
    add_body_text(doc, 'A single income stream is a single point of failure. Economies shift, companies lay off, industries disappear. A modest side-hustle earning just £200-500/month adds up to £2,400-6,000/year of financial buffer, accelerated savings, and optionality you didn\'t have before.')
    add_highlight_box(doc, 'One income is a risk.\nTwo is a plan.\nThree is a fortress.')
    doc.add_page_break()

    add_styled_heading(doc, '2. The 7 Side-Hustle Categories (with Real Earnings)', level=1)
    add_branded_table(doc,
        headers=['Category', 'Setup Time', 'Realistic Monthly £'],
        data=[
            ('1. Freelance services',   '1-2 weeks',  '£200 – £3,000'),
            ('2. Surveys & user testing','1 week',    '£100 – £400'),
            ('3. E-commerce / reselling','2-4 weeks', '£100 – £2,000'),
            ('4. Tutoring / teaching',  '1-2 weeks',  '£200 – £1,500'),
            ('5. Content creation',     '3-12 months','£0 – £5,000+'),
            ('6. Gig economy',          '1 week',     '£300 – £2,000'),
            ('7. Digital products',     '1-3 months', '£50 – £3,000'),
        ])
    doc.add_page_break()

    add_styled_heading(doc, '3. How to Pick YOUR Side-Hustle in 15 Minutes', level=1)
    add_body_text(doc, 'Avoid analysis paralysis. Answer these three questions honestly:')
    _numbers(doc, [
        'What can you do TODAY for someone without training? (Writing, driving, coding, tutoring, selling.)',
        'How much time can you commit per week? (Under 5 hrs → surveys/gig. 5-10 hrs → freelance/tutoring. 10+ hrs → e-commerce/content.)',
        'Do you want fast cash (£100-500 this month) or long-term income (£1,000+ in 6 months)?',
    ])
    add_body_text(doc, 'Match your answers to the table above. Pick ONE. You can always add a second once the first is paying.')
    doc.add_page_break()

    add_styled_heading(doc, '4. The 30-Day Launch Plan', level=1)
    add_styled_heading(doc, 'Week 1 — Research & Register', level=2)
    _bullets(doc, [
        'Pick your hustle and ONE platform to start on (Fiverr, Upwork, Prolific, Vinted, etc.).',
        'Create a professional profile: clear photo, concise bio, specific skill focus.',
        'Set up a separate PayPal/bank account for business income.',
    ])
    add_styled_heading(doc, 'Week 2 — Build the Minimum Offer', level=2)
    _bullets(doc, [
        'List 3 services or 10 products — whatever applies.',
        'Price at the BOTTOM of the market — first reviews matter more than first profits.',
        'Write one sample piece / take product photos / record intro video.',
    ])
    add_styled_heading(doc, 'Week 3 — Launch & Apply', level=2)
    _bullets(doc, [
        'Apply to 10 jobs/day on freelance platforms.',
        'Share your shop on Facebook, LinkedIn, Instagram.',
        'Ask 3 friends/family to leave genuine reviews or buy first.',
    ])
    add_styled_heading(doc, 'Week 4 — Optimise', level=2)
    _bullets(doc, [
        'Analyse what got responses vs what didn\'t. Double down on what works.',
        'Raise prices slightly on any successful listing.',
        'Get your first paying customer this week — that\'s the only real goal.',
    ])
    doc.add_page_break()

    add_styled_heading(doc, '5. Finding Your First 3 Customers', level=1)
    add_body_text(doc, 'First customers are the hardest. Try these in order:')
    _numbers(doc, [
        'Your existing network — tell 20 people you know what you\'re launching.',
        'Facebook groups in your niche — answer questions, offer help, mention your service when relevant.',
        'LinkedIn — post 3x/week about your expertise.',
        'Freelance platforms — apply to 5 jobs/day for 2 weeks.',
        'Cold email — 20 targeted prospects per week with a specific, tailored offer.',
    ])
    add_highlight_box(doc, 'Your first 3 customers probably come from people\nyou already know. Start there before going cold.')
    doc.add_page_break()

    add_styled_heading(doc, '6. Common Traps to Avoid', level=1)
    _bullets(doc, [
        'Buying courses instead of starting. You\'ll learn 10× faster from first customers than any course.',
        'Perfecting the logo/website before anyone has paid you. These DON\'T matter in month 1.',
        'Pricing too high, too early. Get reviews first, then raise prices monthly.',
        'Working for free beyond the first few samples. Free work attracts free-work clients.',
        'Switching hustles after 2 weeks. Pick one. Give it 90 days minimum.',
        'Forgetting tax — track every £ from day 1 (see our UK Tax Basics guide).',
    ])
    doc.add_page_break()

    add_styled_heading(doc, '7. Scaling From £200 to £2,000 a Month', level=1)
    add_body_text(doc, 'Once you\'re earning £200/month, scaling follows a predictable path:')
    _numbers(doc, [
        'Raise prices by 20% every 30 days until customers push back.',
        'Turn down the lowest-paying 20% of work each month — replace with higher-paying clients.',
        'Productise — turn bespoke services into fixed packages (3 tiers: basic / standard / premium).',
        'Build systems — templates, checklists, email sequences — so you spend less time per customer.',
        'When you\'re fully booked, add a waiting list + raise prices 30%.',
    ])
    add_highlight_box(doc, '£200/month grows to £2,000 month about 6-12 months\nfor anyone who keeps raising prices and saying no.')
    doc.add_page_break()

    add_styled_heading(doc, '8. Your Action Plan', level=1)
    _numbers(doc, [
        'Day 1: Pick ONE side-hustle category. No second-guessing.',
        'Day 2: Open a profile on the single best platform for it.',
        'Day 3: Build your minimum offer (3 services or 10 products).',
        'Day 4-7: Apply / list / post every single day.',
        'Week 2: Land first customer, even unpaid. Get a review.',
        'Week 3: Land first PAID customer. Celebrate.',
        'Week 4: Raise prices, ask for more testimonials, re-invest earnings into better tools.',
    ])
    add_body_text(doc, 'Visit www.incomeonline.info — the directory lists 199+ verified platforms across all 7 side-hustle categories, ranked by earning potential and UK availability.')
    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    os.makedirs('/app/backend/static', exist_ok=True)
    for fn, out in [
        (generate_credit_score_document,   'UK_Credit_Score_Masterclass.docx'),
        (generate_isa_vs_sipp_document,    'ISA_vs_SIPP_Complete_Guide.docx'),
        (generate_side_hustle_document,    'Side_Hustle_Quick_Start_Guide.docx'),
    ]:
        with open(f'/app/backend/static/{out}', 'wb') as f:
            f.write(fn().read())
        print(f'Generated {out}')
