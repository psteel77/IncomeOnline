"""Generate the Emergency Fund Guide using the MoneyRules template."""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, PURPLE, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_emergency_fund_document():
    doc = create_moneyrules_document(
        title='Build a 3-Month Emergency Fund',
        subtitle='The Single Best Money Move You Can Make This Year'
    )
    add_title_page(doc,
        title='The Emergency Fund',
        subtitle='How to Build 3 Months of Security — Fast',
        tagline='Why 3 months matters, where to keep it,\nand how to get there in 6–18 months flat.')

    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()
    for num, t in [
        ('1.', 'Introduction — Why an Emergency Fund Changes Your Life'),
        ('2.', 'Why Exactly Three Months?'),
        ('3.', 'Calculating YOUR Number'),
        ('4.', 'Where to Keep It (and Where NOT to)'),
        ('5.', 'The Starter Fund: Your First £1,000'),
        ('6.', 'Scaling to 3 Months — Four Speed Tiers'),
        ('7.', 'What Counts as a Real Emergency?'),
        ('8.', 'Replenishing After You Use It'),
        ('9.', 'Common Mistakes and How to Avoid Them'),
        ('10.', 'Key Takeaways and Action Plan'),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f'{num}  '); r.font.size = Pt(13); r.font.color.rgb = PURPLE; r.bold = True
        r2 = p.add_run(t); r2.font.size = Pt(13); r2.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run('DISCLAIMER: Educational content only. Not financial advice.')
    run.font.size = Pt(9); run.font.color.rgb = GREY; run.italic = True
    doc.add_page_break()

    add_styled_heading(doc, '1. Introduction — Why an Emergency Fund Changes Your Life', level=1)
    add_body_text(doc, 'Ask any financially secure person what their single most valuable money habit is, and the answer is almost always the same: a well-stocked emergency fund. Not a fancy investment. Not a clever tax loophole. Just cash — set aside, untouchable, ready.')
    add_body_text(doc, 'When you have 3 months of essential expenses sitting in a savings account, something profound shifts. A broken boiler is an inconvenience, not a crisis. A redundancy is stressful, but survivable. An unexpected medical bill gets paid in cash, not on a credit card at 29% APR.')
    add_highlight_box(doc, 'An emergency fund doesn\'t make you rich.\nIt makes you unshakeable.')
    doc.add_page_break()

    add_styled_heading(doc, '2. Why Exactly Three Months?', level=1)
    add_body_text(doc, 'You\'ll see every figure from 1 to 12 months recommended online. Here\'s why 3 is the practical sweet spot for most people:')
    add_branded_table(doc,
        headers=['Fund Size', 'Protects Against', 'Opportunity Cost'],
        data=[
            ('1 month',   'Small shocks (boiler, car, vet)',        'Low — often inadequate'),
            ('3 months',  'Most job losses, major repairs',         'Low — the sweet spot'),
            ('6 months',  'Extended unemployment in your sector',   'Medium — large cash drag'),
            ('12 months', 'Career changes, chronic illness risk',   'High — better to invest some'),
        ])
    add_body_text(doc, 'For someone in a stable job with in-demand skills, 3 months comfortably covers the average UK unemployment spell (around 8–12 weeks in 2026). For someone self-employed or in a volatile industry, bump it to 6.')
    add_highlight_box(doc, 'Three months is enough to handle 90% of real-life shocks\nwithout over-hoarding cash you could be investing.')
    doc.add_page_break()

    add_styled_heading(doc, '3. Calculating YOUR Number', level=1)
    add_body_text(doc, 'DO NOT use 3 months of gross pay. That inflates the target badly. The right number is 3 months of essential expenses — the minimum bills you would need to pay if you suddenly had no income.')
    add_styled_heading(doc, 'What to Include', level=2)
    for t in [
        'Rent or mortgage',
        'Council tax, utilities (gas, electric, water, basic internet)',
        'Groceries (cooking at home)',
        'Essential transport (bus/train pass, essential petrol, car insurance)',
        'Minimum debt payments',
        'Childcare',
        'Health insurance / prescriptions',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_styled_heading(doc, 'What to EXCLUDE', level=2)
    for t in [
        'Eating out, takeaways, coffee shops',
        'Subscriptions you could pause (Netflix, Spotify, gym)',
        'Hobby spending',
        'Holidays',
        'Clothes beyond absolute essentials',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_highlight_box(doc, 'Monthly essentials × 3 = Your target.\nFor most UK households this is £3,000 – £6,000.')
    doc.add_page_break()

    add_styled_heading(doc, '4. Where to Keep It (and Where NOT to)', level=1)
    add_styled_heading(doc, 'WHERE: Easy-Access Savings or Cash ISA', level=2)
    add_body_text(doc, 'The whole point of an emergency fund is availability. You need the money within 24 hours, no ifs, no penalties. An easy-access savings account or cash ISA from any reputable UK bank does this. In 2026, rates of 4.5–5% AER are readily available.')
    add_styled_heading(doc, 'WHERE NOT: Fixed Accounts, Stocks, Crypto', level=2)
    add_body_text(doc, 'Never keep an emergency fund in anything that could be down 20% when you need it. Stocks drop. Crypto crashes. Fixed-term bonds lock your money up for months. The slightly higher return is not worth the risk of being unable to access it during an actual emergency.')
    add_branded_table(doc,
        headers=['Location', 'Access Time', 'Safety', 'Verdict'],
        data=[
            ('Easy-access savings',  'Instant',  'FSCS protected',     'YES — ideal'),
            ('Cash ISA',             'Instant',  'FSCS protected',     'YES — tax-free'),
            ('Premium Bonds',        '2–3 days', 'HM Govt backed',     'OK — small delay'),
            ('Fixed 1-year savings', '12 months locked', 'Protected',  'NO — not liquid'),
            ('Stocks & Shares ISA',  'Instant',  'Market risk',        'NO — can crash'),
            ('Current account',      'Instant',  'No interest',        'PARTIAL — only for £500 buffer'),
        ])
    doc.add_page_break()

    add_styled_heading(doc, '5. The Starter Fund: Your First £1,000', level=1)
    add_body_text(doc, 'Before anything else — before debt repayment, before investing, before holidays — build a £500–£1,000 starter fund. This is NOT your full emergency fund. It\'s a buffer that absorbs the small shocks that would otherwise force you back into credit card debt.')
    add_styled_heading(doc, 'How to Build £1,000 in 8 Weeks', level=2)
    for t in [
        'Cancel every non-essential subscription for 2 months (Netflix, Spotify, gym, Amazon Prime).',
        'Sell 10 things on eBay, Vinted or Facebook Marketplace.',
        'Pick up one week of overtime, extra shifts, or a freelance gig via Upwork/Fiverr.',
        'Automate a £200 weekly transfer on payday into a new savings account.',
        'Skip all discretionary dining for 2 months — cook every meal.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_highlight_box(doc, 'The £1,000 starter fund is the\nsingle highest-ROI money habit for beginners.')
    doc.add_page_break()

    add_styled_heading(doc, '6. Scaling to 3 Months — Four Speed Tiers', level=1)
    add_body_text(doc, 'Once your £1,000 starter is in place, here\'s how long the full 3-month target takes at different savings rates (assuming £4,500 target = £1,500/month essentials):')
    add_branded_table(doc,
        headers=['Monthly Contribution', 'Time from £1,000 to £4,500', 'Effort Level'],
        data=[
            ('£150/month',   '24 months',   'Very easy — set & forget'),
            ('£300/month',   '12 months',   'Easy — minor lifestyle trim'),
            ('£500/month',   '7 months',    'Medium — serious budgeting'),
            ('£750/month',   '5 months',    'High — aggressive cuts + side income'),
        ])
    add_body_text(doc, 'For most households, 12–18 months is the realistic target. Faster is better, but SUSTAINABLE is everything — an emergency fund you built in 4 months of misery will often get spent in week 1 on a "reward" holiday. Build it gradually.')
    doc.add_page_break()

    add_styled_heading(doc, '7. What Counts as a Real Emergency?', level=1)
    add_body_text(doc, 'The fund is only as useful as your discipline about when to tap it. Use this test: an emergency is something that is (a) unexpected, (b) urgent, AND (c) essential.')
    add_branded_table(doc,
        headers=['Scenario', 'Emergency?'],
        data=[
            ('Boiler breaks in January',             'YES (unexpected + urgent + essential)'),
            ('Redundancy',                            'YES'),
            ('Emergency vet bill for pet',            'YES'),
            ('A cracked tooth requiring extraction',  'YES'),
            ('Sale at your favourite shop',           'NO'),
            ('Holiday you want but didn\'t save for', 'NO'),
            ('Christmas presents',                    'NO (predictable)'),
            ('Car MOT / service',                     'NO (predictable)'),
        ],
        header_color='DB2777')
    add_highlight_box(doc, 'If it\'s predictable, it\'s a BUDGET item.\nIf it\'s unexpected + urgent + essential, it\'s an EMERGENCY.')
    doc.add_page_break()

    add_styled_heading(doc, '8. Replenishing After You Use It', level=1)
    add_body_text(doc, 'Using your emergency fund is not a failure — it\'s exactly what it exists for. The critical habit is treating replenishment as non-negotiable the moment the crisis passes.')
    for t in [
        'Week 1 after crisis: pause all discretionary spending.',
        'Week 2: set up an automatic monthly contribution to rebuild.',
        'Month 1–6: treat the fund as priority #2 after minimum debt payments.',
        'Do NOT invest new money until the fund is back to full strength.',
        'Do NOT delay because "I\'ll top it up next year". Next year becomes never.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    doc.add_page_break()

    add_styled_heading(doc, '9. Common Mistakes and How to Avoid Them', level=1)
    for head, body in [
        ('Keeping it in your current account',
         'The money must be mentally AND physically separate. If you can see it alongside your spending money, you will spend it.'),
        ('"Investing" it for higher returns',
         'Stock market funds drop 30%+ in crashes. The day you need the money will often be the day the market is down. Keep this cash in cash.'),
        ('Setting the target too high',
         'Aiming for 12 months before even starting is paralysis. Build 1 month first. Then 2. Then 3. Never let perfect be the enemy of done.'),
        ('Raiding it for non-emergencies',
         'Birthdays and Christmas are not emergencies. They are scheduled. Budget for them separately.'),
        ('Forgetting inflation',
         'Review your target number once a year. £4,500 in 2020 probably needs to be £5,500 in 2026 for the same standard of living.'),
    ]:
        add_styled_heading(doc, head, level=2)
        add_body_text(doc, body)
    doc.add_page_break()

    add_styled_heading(doc, '10. Key Takeaways and Action Plan', level=1)
    for t in [
        '3 months of ESSENTIAL expenses is the target (not 3 months of gross pay).',
        'For most UK households this works out at £3,000–£6,000.',
        'Keep it in easy-access savings or a cash ISA. Never in stocks or crypto.',
        'Build £500–£1,000 as a starter buffer FIRST, before the full fund.',
        'Automate contributions — willpower fails, direct debits don\'t.',
        'Use it only for unexpected + urgent + essential events.',
        'Replenish aggressively after any use.',
    ]:
        p = doc.add_paragraph(t, style='List Bullet')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT

    add_styled_heading(doc, 'Your First Week', level=2)
    for t in [
        'Day 1: Calculate your essential monthly expenses (cooking + bills only).',
        'Day 2: Multiply by 3. That\'s your target.',
        'Day 3: Open a separate easy-access savings account.',
        'Day 4: Transfer whatever you can afford today, even £20.',
        'Day 5: Set up an automatic £200 monthly direct debit on payday.',
        'Day 6: Cancel 3 subscriptions you don\'t truly need and redirect to the fund.',
        'Day 7: Diarise a review for 3 months time.',
    ]:
        p = doc.add_paragraph(t, style='List Number')
        for r in p.runs: r.font.size = Pt(11); r.font.color.rgb = BODY_TEXT
    add_body_text(doc, 'Visit www.incomeonline.info to discover 199+ platforms that can accelerate your savings rate — even an extra £100/month cuts your timeline nearly in half.')
    add_closing_page(doc)
    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    buf = generate_emergency_fund_document()
    os.makedirs('/app/backend/static', exist_ok=True)
    with open('/app/backend/static/The_Emergency_Fund_Guide.docx', 'wb') as f:
        f.write(buf.read())
    print('Emergency Fund guide generated!')
