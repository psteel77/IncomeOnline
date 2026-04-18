"""
Generate The 50/30/20 Budget Rule guide using the MoneyRules branded template.
"""
from moneyrules_template import (
    create_moneyrules_document, add_title_page, add_styled_heading,
    add_body_text, add_highlight_box, add_branded_table, add_closing_page,
    save_to_buffer, DEEP_PURPLE, PURPLE, PINK, ACCENT_GOLD, DARK_TEXT, BODY_TEXT, GREY
)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_503020_document():
    """Generate the complete 50/30/20 Budget Rule Word document."""
    doc = create_moneyrules_document(
        title='The 50/30/20 Budget Rule',
        subtitle='The Simplest Way to Take Control of Your Money'
    )

    # ===== TITLE PAGE =====
    add_title_page(
        doc,
        title='The 50/30/20 Rule',
        subtitle='A Complete Guide to Budgeting Your Income',
        tagline='Split every pound you earn into Needs, Wants and Savings\n— the three-bucket system that actually works.'
    )

    # ===== TABLE OF CONTENTS =====
    add_styled_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph()

    toc_items = [
        ('1.', 'Introduction — Why Most Budgets Fail'),
        ('2.', 'The 50/30/20 Rule Explained'),
        ('3.', 'Category One: Needs (50%)'),
        ('4.', 'Category Two: Wants (30%)'),
        ('5.', 'Category Three: Savings & Debt (20%)'),
        ('6.', 'Step-by-Step: Build Your 50/30/20 Budget'),
        ('7.', 'Worked Examples at Different Income Levels'),
        ('8.', 'Common Mistakes and How to Avoid Them'),
        ('9.', 'Adapting the Rule to Your Situation'),
        ('10.', 'Key Takeaways and Next Steps'),
    ]

    for num, title_text in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(10)
        run_num = para.add_run(f'{num}  ')
        run_num.font.size = Pt(13)
        run_num.font.color.rgb = PURPLE
        run_num.bold = True
        run_title = para.add_run(title_text)
        run_title.font.size = Pt(13)
        run_title.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run(
        'DISCLAIMER: This document is for educational and informational purposes only. '
        'It does not constitute financial advice. Always consult a qualified financial '
        'adviser before making significant financial decisions.'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    run.italic = True
    doc.add_page_break()

    # ===== CHAPTER 1: INTRODUCTION =====
    add_styled_heading(doc, '1. Introduction — Why Most Budgets Fail', level=1)

    add_body_text(doc,
        'If you have ever tried to budget, you already know the feeling: you download a '
        'spreadsheet with forty-seven categories, spend a Sunday afternoon filling it in, '
        'and by Wednesday you have stopped logging receipts. Two weeks later, the '
        'spreadsheet is closed forever.')

    add_body_text(doc,
        'Most budgeting systems fail for exactly the same reason — they demand too much '
        'precision for too little reward. You do not need to track whether your coffee '
        'spend was £3.20 or £3.40; you need to know whether, overall, you are living '
        'within your means, enjoying your life, and building financial security.')

    add_body_text(doc,
        'That is the genius of the 50/30/20 Rule. Popularised by US Senator Elizabeth '
        'Warren and her daughter Amelia Warren Tyagi in their 2005 book "All Your Worth", '
        'it replaces dozens of line items with just three simple buckets — and it works '
        'for almost anyone, regardless of income level.')

    add_highlight_box(doc,
        'The 50/30/20 Rule gives you a budget that fits on a sticky note\n'
        '— and still covers every pound you earn.')

    add_styled_heading(doc, 'What This Guide Will Teach You', level=2)
    for item in [
        'What the 50/30/20 Rule is and why it works so well.',
        'How to categorise every expense into Needs, Wants, or Savings.',
        'How to build your own budget in under 30 minutes.',
        'Real examples at £20k, £35k and £60k income levels.',
        'How to adapt the rule if you live in a high-rent city or have heavy debt.',
    ]:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT
    doc.add_page_break()

    # ===== CHAPTER 2: THE RULE =====
    add_styled_heading(doc, '2. The 50/30/20 Rule Explained', level=1)

    add_body_text(doc,
        'The rule takes your after-tax income (the money that actually lands in your bank '
        'account each month) and splits it into three fixed percentages:')

    add_branded_table(doc,
        headers=['Bucket', 'Percentage', 'Purpose'],
        data=[
            ('Needs',         '50%', 'Essential living expenses you cannot avoid'),
            ('Wants',         '30%', 'Lifestyle spending that makes life enjoyable'),
            ('Savings & Debt', '20%', 'Future-you money: savings, investments, extra debt payments'),
        ]
    )

    add_body_text(doc,
        'That is the entire framework. No forty-seven categories, no guilt-tripping '
        'over a cup of coffee, no colour-coded spreadsheet. Three buckets, three '
        'percentages, one simple goal: make sure every pound has a job.')

    add_styled_heading(doc, 'Why Percentages Work Better Than Fixed Amounts', level=2)
    add_body_text(doc,
        'Traditional budgets set fixed amounts ("£400 for groceries, £200 for petrol"). '
        'The problem is that life is not fixed. You get a pay rise; your rent goes up; '
        'inflation hits. Fixed-amount budgets break the moment anything changes.')

    add_body_text(doc,
        'Percentage-based budgets self-adjust. If you earn more, all three buckets grow '
        'proportionally. If you earn less, they all shrink together. The discipline '
        'stays the same, and the maths keeps working forever.')

    add_highlight_box(doc,
        'Fixed-amount budgets break at the first pay rise.\n'
        'Percentage budgets last a lifetime.')
    doc.add_page_break()

    # ===== CHAPTER 3: NEEDS =====
    add_styled_heading(doc, '3. Category One: Needs (50%)', level=1)

    add_body_text(doc,
        'Needs are expenses you genuinely cannot avoid without serious consequences. '
        'If you stopped paying them, you would lose your home, your job, your health, '
        'or your ability to function day-to-day.')

    add_body_text(doc,
        'That sounds obvious — but here is where most people go wrong. "Need" is not '
        'the same as "I pay for it every month". Gym memberships, Netflix, Spotify and '
        'takeaway coffees are NOT needs, no matter how regular they are.')

    add_styled_heading(doc, 'What Counts as a Need', level=2)

    add_branded_table(doc,
        headers=['Category', 'Examples'],
        data=[
            ('Housing',        'Rent or mortgage payment, council tax, building insurance'),
            ('Utilities',      'Gas, electric, water, internet, basic phone bill'),
            ('Groceries',      'Actual food you cook at home (not restaurant meals)'),
            ('Transport',      'Bus pass, essential petrol, car insurance, MOT'),
            ('Insurance',      'Health, life, contents insurance'),
            ('Minimum Debt',   'Minimum required payments on loans and credit cards'),
            ('Childcare',      'Nursery fees, essential school costs'),
        ],
        header_color='DB2777'
    )

    add_styled_heading(doc, 'The Honest Test', level=2)
    add_body_text(doc,
        'Ask yourself: "If I stopped paying this for three months, would something '
        'genuinely bad happen?" If the honest answer is no, it belongs in Wants.')

    add_highlight_box(doc,
        'If your Needs exceed 50%, you have a lifestyle-cost problem,\n'
        'not a budgeting problem. Time to look at the big levers.')
    doc.add_page_break()

    # ===== CHAPTER 4: WANTS =====
    add_styled_heading(doc, '4. Category Two: Wants (30%)', level=1)

    add_body_text(doc,
        'Wants are everything that makes life enjoyable — the things you choose to spend '
        'money on because they bring pleasure, convenience or entertainment, even though '
        'you could technically live without them.')

    add_body_text(doc,
        'This is the most misunderstood part of the 50/30/20 Rule. Many people hear '
        '"budget" and assume it means cutting out fun. The truth is the opposite: this '
        'rule formally protects 30% of your income for enjoyment. You are allowed — in '
        'fact, required — to spend it.')

    add_styled_heading(doc, 'What Counts as a Want', level=2)

    add_branded_table(doc,
        headers=['Category', 'Examples'],
        data=[
            ('Eating Out',      'Restaurants, takeaways, coffee shops, lunches at work'),
            ('Entertainment',   'Netflix, Spotify, cinema, concerts, gigs, gaming'),
            ('Hobbies',         'Gym membership, sports, crafts, collections'),
            ('Shopping',        'Clothes, gadgets, home decor (beyond essentials)'),
            ('Travel',          'Holidays, weekend breaks, day trips'),
            ('Upgrades',        'Faster internet, nicer car, premium phone, designer brands'),
            ('Gifts',           'Birthdays, Christmas, weddings, generosity'),
        ],
        header_color='DB2777'
    )

    add_styled_heading(doc, 'The Upgrade Principle', level=2)
    add_body_text(doc,
        'One subtle point: if your phone contract is £15/month, that part is a Need '
        '(basic communication). But if you upgrade to a £60/month contract with a '
        'flagship handset, the extra £45 is a Want. Always separate the "basic" version '
        'from the "upgraded" version — the upgrade goes in Wants.')

    add_highlight_box(doc,
        'Wants are not the enemy of a good budget.\n'
        'They are the reason you work.')
    doc.add_page_break()

    # ===== CHAPTER 5: SAVINGS & DEBT =====
    add_styled_heading(doc, '5. Category Three: Savings & Debt (20%)', level=1)

    add_body_text(doc,
        'This is the bucket that transforms your financial life. Twenty pence of every '
        'pound you earn goes directly towards your future self — through savings, '
        'investments, or accelerated debt repayment.')

    add_body_text(doc,
        'Most people get this backwards. They pay every bill, spend what feels right on '
        'themselves, and save whatever happens to be left over (usually nothing). The '
        '50/30/20 Rule flips this: you pay your future first, and spend what is left.')

    add_styled_heading(doc, 'What Goes in This Bucket', level=2)

    add_branded_table(doc,
        headers=['Priority', 'Destination'],
        data=[
            ('1 (first)',   'Emergency fund until you have 3-6 months of Needs covered'),
            ('2',           'High-interest debt paid down aggressively (credit cards first)'),
            ('3',           'Workplace pension — at least up to the full employer match'),
            ('4',           'Stocks & Shares ISA or SIPP for long-term investing'),
            ('5',           'Specific goals — house deposit, wedding, career break'),
        ],
        header_color='EA580C'
    )

    add_styled_heading(doc, 'The Golden Order', level=2)
    for item in [
        'Build a starter emergency fund of £1,000 before anything else.',
        'Claim your full employer pension match — it is literally free money.',
        'Crush any debt with an interest rate above 8%.',
        'Grow your emergency fund to 3-6 months of essential expenses.',
        'Max out ISA / pension contributions for long-term compounding.',
    ]:
        para = doc.add_paragraph(item, style='List Number')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    add_highlight_box(doc,
        'Pay your future self FIRST — on the day your salary arrives.\n'
        'Automate the transfer before a single pound is spent.')
    doc.add_page_break()

    # ===== CHAPTER 6: STEP-BY-STEP =====
    add_styled_heading(doc, '6. Step-by-Step: Build Your 50/30/20 Budget', level=1)

    add_body_text(doc, 'Here is exactly how to set up your budget in under 30 minutes.')

    add_styled_heading(doc, 'Step 1: Calculate Your Monthly Take-Home Pay', level=2)
    add_body_text(doc,
        'Look at your last three payslips and find the net figure — the amount that '
        'actually lands in your bank account after tax, National Insurance and pension '
        'contributions. Use the average of the three as your monthly income.')

    add_styled_heading(doc, 'Step 2: Calculate Your Target Amounts', level=2)
    add_body_text(doc, 'Multiply your monthly take-home pay by the three percentages:')

    add_branded_table(doc,
        headers=['Bucket', 'Calculation', 'Example (£2,500/mo)'],
        data=[
            ('Needs',          'Monthly pay × 0.50',  '£1,250'),
            ('Wants',          'Monthly pay × 0.30',  '£750'),
            ('Savings & Debt', 'Monthly pay × 0.20',  '£500'),
        ]
    )

    add_styled_heading(doc, 'Step 3: Audit Your Last 3 Months', level=2)
    add_body_text(doc,
        'Download your last three bank statements. Tag every single outgoing transaction '
        'as N (Need), W (Want) or S (Savings). Add up each category and divide by three '
        'to get your current monthly averages.')

    add_styled_heading(doc, 'Step 4: Compare to Target and Adjust', level=2)
    add_body_text(doc,
        'Now compare your current split to the 50/30/20 target. Wherever you are over '
        'budget, identify one or two changes you can make this month. Do not try to fix '
        'everything at once — aim for 5% shifts at a time.')

    add_styled_heading(doc, 'Step 5: Automate Everything', level=2)
    add_body_text(doc,
        'On payday, set up three automatic standing orders: one to a "Bills" account '
        '(Needs), one to a "Spending" account (Wants), and one to a Savings/ISA account '
        '(the 20% bucket). Once automated, the budget runs itself.')

    add_highlight_box(doc,
        'A budget that requires willpower every day will fail.\n'
        'A budget that runs on automatic standing orders will last forever.')
    doc.add_page_break()

    # ===== CHAPTER 7: WORKED EXAMPLES =====
    add_styled_heading(doc, '7. Worked Examples at Different Income Levels', level=1)

    add_body_text(doc,
        'The beauty of the percentage-based approach is that it scales to any income. '
        'Here are three worked examples.')

    add_styled_heading(doc, 'Example 1: £20,000 Gross Salary (~£1,550/mo net)', level=2)

    add_branded_table(doc,
        headers=['Bucket', 'Amount', 'Typical Allocation'],
        data=[
            ('Needs (50%)',          '£775', 'Rent £500, bills £150, groceries £100, transport £25'),
            ('Wants (30%)',          '£465', 'Eating out £150, subscriptions £40, social £150, clothes £125'),
            ('Savings & Debt (20%)', '£310', '£200 to ISA, £110 to emergency fund'),
        ]
    )

    add_body_text(doc,
        'At this income, living in a low-cost area or house-sharing is often essential '
        'to keep Needs at or below 50%. It is a tight budget, but achievable.')

    add_styled_heading(doc, 'Example 2: £35,000 Gross Salary (~£2,400/mo net)', level=2)

    add_branded_table(doc,
        headers=['Bucket', 'Amount', 'Typical Allocation'],
        data=[
            ('Needs (50%)',          '£1,200', 'Rent £800, bills £200, groceries £150, transport £50'),
            ('Wants (30%)',          '£720',   'Eating out £200, subs £60, social £200, hobbies £260'),
            ('Savings & Debt (20%)', '£480',   '£300 to ISA, £150 to pension top-up, £30 to emergency'),
        ],
        header_color='DB2777'
    )

    add_body_text(doc,
        'At this income, the 50/30/20 Rule starts to feel genuinely comfortable — '
        'enough for a proper life and still meaningful long-term savings.')

    add_styled_heading(doc, 'Example 3: £60,000 Gross Salary (~£3,750/mo net)', level=2)

    add_branded_table(doc,
        headers=['Bucket', 'Amount', 'Typical Allocation'],
        data=[
            ('Needs (50%)',          '£1,875', 'Mortgage £1,100, bills £275, groceries £250, transport £250'),
            ('Wants (30%)',          '£1,125', 'Dining £300, travel £400, hobbies £250, shopping £175'),
            ('Savings & Debt (20%)', '£750',   '£500 ISA, £200 pension boost, £50 specific goals'),
        ],
        header_color='EA580C'
    )

    add_body_text(doc,
        'At this level many people choose to increase savings beyond 20% — sometimes '
        'called "50/20/30" — pushing more into long-term investments while their '
        'lifestyle is still modest. This is how financial independence is built.')
    doc.add_page_break()

    # ===== CHAPTER 8: MISTAKES =====
    add_styled_heading(doc, '8. Common Mistakes and How to Avoid Them', level=1)

    add_styled_heading(doc, 'Mistake 1: Using Gross Income Instead of Net', level=2)
    add_body_text(doc,
        'The rule applies to take-home pay, not your gross salary. If you use gross '
        'income, you will always be over budget because you are trying to spend money '
        'that HMRC already took.')

    add_styled_heading(doc, 'Mistake 2: Hiding Wants Inside Needs', level=2)
    add_body_text(doc,
        'Classifying Netflix as a "utility" or the gym as "health" is the most common '
        'way people cheat the rule. Be ruthlessly honest — if it is not essential for '
        'survival, safety or work, it is a Want.')

    add_styled_heading(doc, 'Mistake 3: Skipping the Audit', level=2)
    add_body_text(doc,
        'Building a budget without auditing three months of real spending is fiction. '
        'You need to know what you ACTUALLY spend before you can plan what you WILL '
        'spend.')

    add_styled_heading(doc, 'Mistake 4: Cutting Wants to Zero', level=2)
    add_body_text(doc,
        'Enthusiastic new budgeters often try to slash Wants to £0 and pour everything '
        'into savings. This always backfires — usually within two months — in a splurge '
        'that wipes out the progress. The 30% is there for a reason: it keeps you sane '
        'and sustainable.')

    add_styled_heading(doc, 'Mistake 5: Budgeting Alone When You Live With a Partner', level=2)
    add_body_text(doc,
        'If you share finances, build the budget together. Money arguments are the '
        'leading cause of relationship stress — a shared 50/30/20 plan on a single '
        'page removes most of those conversations before they happen.')

    add_highlight_box(doc,
        'Honesty is the hardest part of budgeting.\n'
        'Once you are honest, the maths is trivial.')
    doc.add_page_break()

    # ===== CHAPTER 9: ADAPTATION =====
    add_styled_heading(doc, '9. Adapting the Rule to Your Situation', level=1)

    add_body_text(doc,
        'The 50/30/20 Rule is a starting point, not a straightjacket. Here is how to '
        'adapt it to real life.')

    add_styled_heading(doc, 'High Cost of Living (e.g. London, Edinburgh, Manchester)', level=2)
    add_body_text(doc,
        'If rent alone eats 40% of your income, getting Needs down to 50% is genuinely '
        'hard. A fair adaptation is 60/20/20 — accepting higher Needs but protecting '
        'the 20% savings rate. Do NOT let Wants shrink to zero; cut Needs instead by '
        'house-sharing, relocating, or negotiating a raise.')

    add_styled_heading(doc, 'High-Debt Situation', level=2)
    add_body_text(doc,
        'If you have credit card debt above 15% APR, flip the rule temporarily to '
        '50/20/30 — with 30% going to debt destruction. This gets you out of the '
        'trap faster and is usually better than any investment return you could earn.')

    add_styled_heading(doc, 'High Savings Rate Goal (FIRE Movement)', level=2)
    add_body_text(doc,
        'For those pursuing Financial Independence, 50/20/30 or even 50/10/40 are '
        'common variations. The higher the savings rate, the faster you reach '
        'financial independence — but your quality of life today matters too.')

    add_branded_table(doc,
        headers=['Situation', 'Suggested Split', 'Priority'],
        data=[
            ('Standard / balanced',     '50 / 30 / 20', 'Long-term wealth'),
            ('High cost of living',     '60 / 20 / 20', 'Preserve savings rate'),
            ('Aggressive debt payoff',  '50 / 20 / 30', 'Eliminate high-interest debt'),
            ('Early retirement (FIRE)', '50 / 10 / 40', 'Maximum savings rate'),
            ('Low income / survival',   '70 / 20 / 10', 'Build basic emergency fund'),
        ],
        header_color='DB2777'
    )

    add_highlight_box(doc,
        'The exact percentages matter less than the habit.\n'
        'What matters is that every pound has a job — on purpose.')
    doc.add_page_break()

    # ===== CHAPTER 10: CONCLUSION =====
    add_styled_heading(doc, '10. Key Takeaways and Next Steps', level=1)

    add_body_text(doc,
        'The 50/30/20 Rule is not the most sophisticated budgeting method in the world. '
        'It is not the most precise. It is not the most optimised. What it is, however, '
        'is the budgeting method most people actually stick with — and a budget you '
        'stick with is worth a thousand perfect budgets you abandon.')

    for t in [
        'Split your net monthly income: 50% Needs, 30% Wants, 20% Savings & Debt.',
        'Needs are genuine essentials. Wants are everything else you enjoy. Savings is future-you.',
        'Pay your future self FIRST, on payday, via an automated standing order.',
        'Audit your last three months of spending before designing your budget.',
        'Use three separate bank accounts to make the system automatic and foolproof.',
        'Adapt the percentages to your situation — high cost of living, high debt, or high savings goals.',
        'Honesty about Needs vs Wants is more important than any percentage.',
        'Review once a quarter. Adjust when life changes. Let the rule run the rest of the time.',
    ]:
        para = doc.add_paragraph(t, style='List Bullet')
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    add_styled_heading(doc, 'Your Action Plan for This Week', level=2)

    add_body_text(doc,
        'Do not wait for a new tax year or a round-number date. This week:')
    for t in [
        'Day 1: Calculate your exact monthly take-home pay (average of last 3 payslips).',
        'Day 2: Download 3 months of bank statements and tag every transaction.',
        'Day 3: Add up Needs, Wants and Savings. Compare to the 50/30/20 target.',
        'Day 4: Identify your TWO biggest Wants and decide if they bring enough joy.',
        'Day 5: Open a separate Savings/ISA account if you do not have one.',
        'Day 6: Set up automatic standing orders for payday.',
        'Day 7: Rest. The system now runs itself.',
    ]:
        para = doc.add_paragraph(t, style='List Number')
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_TEXT

    add_body_text(doc,
        'Visit Income Online at www.incomeonline.info to discover 199+ legitimate '
        'platforms where you can boost the income side of your budget equation. Because '
        'the fastest way to make every percentage bigger is to grow the pie.')

    add_closing_page(doc)

    return save_to_buffer(doc)


if __name__ == '__main__':
    import os
    buf = generate_503020_document()
    os.makedirs('/app/backend/static', exist_ok=True)
    with open('/app/backend/static/The_50_30_20_Budget_Rule.docx', 'wb') as f:
        f.write(buf.read())
    print('50/30/20 document generated successfully!')
